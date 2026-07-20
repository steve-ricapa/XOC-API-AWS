from __future__ import annotations

import os
import tempfile
from datetime import datetime

from src.reports.docx_renderer import (
    generate_report_docx,
    insert_findings_table,
    replace_residual_template_text,
    validate_docx_file,
)
from src.reports.storage import (
    download_artifact,
    download_template,
    template_exists,
    upload_document,
)
from src.shared.logging import logger


def handler(event: dict, context) -> dict:
    document_id = event.get("documentId")
    tenant_id = event.get("tenantId")
    generated_content_key = event.get("generatedContentKey")
    document_type = event.get("documentType", "")

    if not all([document_id, tenant_id, generated_content_key]):
        raise ValueError("documentId, tenantId, and generatedContentKey are required")

    tenant_id = int(tenant_id)

    content = download_artifact(generated_content_key)

    with tempfile.TemporaryDirectory() as tmpdir:
        template_path = os.path.join(tmpdir, "template.docx")
        output_path = os.path.join(tmpdir, "generated.docx")
        chart_path = os.path.join(tmpdir, "chart.png")

        has_template = template_exists(document_type)
        if has_template:
            download_template(document_type, template_path)
        else:
            _create_minimal_template(template_path)

        data = _build_render_data(content)

        try:
            generate_report_docx(template_path, data, output_path, has_real_template=has_template)
        except Exception as exc:
            logger.warning("Template rendering failed: %s. Falling back to direct replacement.", exc)
            from src.reports.docx_renderer import generate_from_real_template
            output_path = generate_from_real_template(template_path, data, output_path)

        insert_findings_table(output_path, data, output_path)
        replace_residual_template_text(output_path, data)

        valid, msg = validate_docx_file(output_path)
        if not valid:
            raise RuntimeError(f"DOCX validation failed: {msg}")

        result = upload_document(tenant_id, document_id, document_type, output_path)
        logger.info("DOCX uploaded for document %s: s3://%s/%s", document_id, result["s3_bucket"], result["s3_key"])

    return {
        **event,
        "docxKey": f"{result['s3_bucket']}/{result['s3_key']}",
        "s3Bucket": result["s3_bucket"],
        "s3Key": result["s3_key"],
        "s3VersionId": result["s3_version_id"],
        "sizeBytes": result["size_bytes"],
    }


def _build_render_data(content: dict) -> dict:
    document = content.get("document", {})
    document_type = content.get("document_type", "minority_report")
    return {
        "report": {
            "id": document.get("id", content.get("sections", [{}])[0].get("id", "document")),
            "title": document.get("title", _default_title(document_type)),
            "service": document.get("service", "Servicio de Generacion Documental XOC"),
            "generated_at": document.get("generated_at", datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S")),
            "period": document.get("period", "Ultima semana evaluada"),
            "prepared_by": document.get("prepared_by", "TXDXSECURE"),
            "executive_summary": document.get("executive_summary", content.get("sections", [{}])[0].get("content", "")),
            "results": document.get("results", "Resultados consolidados del documento."),
        },
        "tenant": {
            "id": "tenant",
            "name": "Cliente",
        },
        "severity_summary": content.get("severity_summary", {}),
        "findings": content.get("findings", []),
        "domains": content.get("domains", []),
        "actions_worked": content.get("actions_worked", []),
        "security_news": content.get("security_news", content.get("sections", [{}])[-1].get("news", [])) if content.get("sections") else [],
        "tools": ["MonEvents", "MonVulE", "MonVulC", "MonApps", "MonNet", "MonInfra"],
    }


def _default_title(document_type: str) -> str:
    if document_type == "small_report":
        return "Small Report - XOC"
    if document_type == "informe_soporte":
        return "Informe de Soporte - XOC"
    return "Minority Report - XOC"


def _create_minimal_template(template_path: str) -> str:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Mm, Pt

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("{{ report.title }}")
    run.bold = True
    run.font.size = Pt(24)

    for text in ["{{ report.service }}", "Cliente: {{ tenant.name }}", "Periodo: {{ report.period }}", "Generado: {{ report.generated_at }}", "Preparado por: {{ report.prepared_by }}"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(text).font.size = Pt(12)

    doc.add_page_break()

    def heading(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(15)

    heading("1. Datos generales")
    doc.add_paragraph("Tenant ID: {{ tenant.id }}")
    doc.add_paragraph("Reporte ID: {{ report.id }}")

    heading("2. Herramientas")
    doc.add_paragraph("{% for tool in tools %}- {{ tool }}{% endfor %}")

    heading("3. Resumen ejecutivo")
    doc.add_paragraph("{{ report.executive_summary }}")

    heading("4. Analisis de severidades")
    doc.add_paragraph("Critico: {{ severity_summary.critical }} | Alto: {{ severity_summary.high }} | Medio: {{ severity_summary.medium }} | Bajo: {{ severity_summary.low }} | Informativo: {{ severity_summary.informational }}")

    heading("5. Grafico de vulnerabilidades")
    doc.add_paragraph("{{ severity_chart }}")

    heading("6. Tabla de findings")
    doc.add_paragraph("{% for finding in findings %}{{ finding.id }} | {{ finding.domain }} | {{ finding.title }} | {{ finding.severity }}\n{% endfor %}")
    doc.add_paragraph("[[FINDINGS_TABLE]]")

    heading("7. Seguridad por dominio")
    doc.add_paragraph("{% for domain in domains %}{{ loop.index }}. {{ domain.name }}\nResumen: {{ domain.summary }}\nFindings relacionados: {{ domain.findings | join(', ') if domain.findings else 'Sin findings asociados' }}\n\n{% endfor %}")

    heading("8. Acciones trabajadas durante la semana")
    doc.add_paragraph("{% for action in actions_worked %}- {{ action }}\n{% endfor %}")

    heading("9. Resultados obtenidos")
    doc.add_paragraph("{{ report.results }}")

    heading("10. Noticias de seguridad")
    doc.add_paragraph("{% for news in security_news %}- {{ news.date }} | {{ news.source }} | {{ news.title }}\n{{ news.summary }}\nEnlaces: {{ news.links | join(', ') }}\n\n{% endfor %}")

    doc.save(template_path)
    return template_path

from __future__ import annotations

import os
import shutil
import tempfile
import zipfile
from copy import deepcopy

from docx import Document
from docx.shared import Inches, Mm, Pt
from docxtpl import DocxTemplate
from jinja2 import Environment, StrictUndefined
from lxml import etree


DOMAIN_TABLE_ORDER = [
    "Dominio de Web Externo",
    "Dominio de IPs Públicas",
    "Dominio de FW",
    "Dominio Infraestructura de Computo - Servers",
    "Dominio Infraestructura de Red - Switches",
    "Dominio Infraestructura de Red - WIFI",
    "Dominio Infraestructura de Computo - Desktops",
    "Dominio Infraestructura OT/IoT",
]


def _set_cell_text(cell, text: str, *, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(10)


def _replace_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def _replace_cover_prepared_line(paragraph, tenant_name: str, prepared_by: str) -> None:
    if len(paragraph.runs) >= 5 and "\t" in paragraph.text:
        paragraph.runs[0].text = tenant_name
        paragraph.runs[1].text = ""
        paragraph.runs[2].text = ""
        paragraph.runs[3].text = "\t"
        paragraph.runs[4].text = prepared_by
        for run in paragraph.runs[5:]:
            run.text = ""
        return
    _replace_paragraph_text(paragraph, f"{tenant_name}\t{prepared_by}")


def _clear_table_rows(table) -> None:
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)


def _add_finding_row(table, record: dict) -> None:
    row = table.add_row().cells
    _set_cell_text(row[0], str(record.get("id", "")))
    _set_cell_text(row[1], str(record.get("title", "")))
    _set_cell_text(row[2], str(record.get("affected_hosts", "")))
    _set_cell_text(row[3], str(record.get("severity", "")))


def _normalize_domain(domain_name: str) -> str:
    mapping = {
        "Dominio de IPs Públicas": "Dominio de IPs Publicas",
        "Dominio de IPs Publicas": "Dominio de IPs Publicas",
        "Dominio Infraestructura de Computo - Desktops": "Dominio Infraestructura de de Computo - Desktops",
    }
    return mapping.get(domain_name, domain_name)


def _group_findings_by_domain(data: dict) -> dict[str, list[dict]]:
    grouped: dict[str, list[dict]] = {}
    for finding in data.get("findings", []):
        grouped.setdefault(_normalize_domain(finding.get("domain", "")), []).append(finding)
    return grouped


def render_from_template(template_path: str, data: dict, output_path: str) -> str:
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    template = DocxTemplate(template_path)
    env = Environment(undefined=StrictUndefined, autoescape=False)

    context = deepcopy(data)
    context.setdefault("severity_chart", "")

    template.render(context, jinja_env=env)
    template.save(output_path)
    return output_path


def replace_known_paragraphs(doc: Document, data: dict) -> None:
    from src.reports.context_builder import build_summary_intro, build_result_paragraphs, build_service_description

    summary_1, summary_2, summary_3 = build_summary_intro(data)
    result_paragraphs = build_result_paragraphs(data)

    if len(doc.paragraphs) >= 169:
        updates = {
            33: data["report"].get("period", ""),
            40: build_service_description(data),
            42: data["report"].get("period", ""),
            50: _build_data_base_paragraph(data),
            53: summary_1,
            54: summary_2,
            56: summary_3,
            58: _build_comparative_text(data),
            62: _build_histogram_text(),
            67: _build_results_summary(),
            69: result_paragraphs[0],
            70: result_paragraphs[1],
            71: result_paragraphs[2],
            72: result_paragraphs[3],
            73: result_paragraphs[4],
            74: result_paragraphs[5],
            76: "Se continuará con la validación de mitigaciones priorizadas, actualización de evidencias y seguimiento de hallazgos críticos y altos.",
            79: "Se programará el seguimiento de hallazgos expuestos y revisión de controles compensatorios aplicados durante el siguiente corte.",
            81: _build_requirement_text(data),
            82: "Se recomienda coordinar la revisión prioritaria de activos con exposición crítica y alta, así como la validación técnica de hallazgos que requieran ventana de mantenimiento.",
            85: data["domains"][0]["summary"],
            91: data["domains"][1]["summary"],
            96: data["domains"][2]["summary"],
            102: data["domains"][3]["summary"],
            108: data["domains"][4]["summary"],
            114: data["domains"][5]["summary"],
            120: data["domains"][6]["summary"],
            126: data["domains"][7]["summary"],
            144: result_paragraphs[0],
            145: result_paragraphs[1],
            146: result_paragraphs[2],
            149: "Queda pendiente la revisión analítica y operativa de los hallazgos de mayor criticidad, así como la confirmación de acciones correctivas implementadas durante el siguiente corte semanal.",
        }
        action_start = 133
        action_slots = 8
        news_paragraph_start = 151
        blocks_per_news = 9
    else:
        updates = {
            34: build_service_description(data),
            37: data["report"].get("period", ""),
            47: _build_data_base_paragraph(data),
            50: summary_1,
            52: summary_2,
            54: summary_3,
            59: _build_comparative_text(data),
            81: _build_histogram_text(),
            87: _build_results_summary(),
            90: result_paragraphs[0],
            92: result_paragraphs[1],
            94: result_paragraphs[2],
            96: result_paragraphs[3],
            98: result_paragraphs[4],
            100: result_paragraphs[5],
            102: data["report"].get("results", ""),
            104: "Se continuará con la validación de mitigaciones priorizadas, actualización de evidencias y seguimiento de hallazgos críticos y altos.",
            106: "Asimismo, se mantendrá el monitoreo sobre activos con exposición recurrente y se actualizarán las recomendaciones conforme al avance de remediación.",
            108: _build_requirement_text(data),
            112: data["domains"][0]["summary"],
            116: data["domains"][1]["summary"],
            123: data["domains"][2]["summary"],
            127: data["domains"][3]["summary"],
            130: data["domains"][4]["summary"],
            133: data["domains"][5]["summary"],
            136: data["domains"][6]["summary"],
            139: data["domains"][7]["summary"],
            158: result_paragraphs[0],
            160: result_paragraphs[1],
            162: result_paragraphs[2],
            164: result_paragraphs[3],
            166: result_paragraphs[4],
            168: result_paragraphs[5],
            172: "Queda pendiente la revisión analítica y operativa de los hallazgos de mayor criticidad, así como la confirmación de acciones correctivas implementadas durante el siguiente corte semanal.",
        }
        action_start = 142
        action_slots = 12
        news_paragraph_start = 175
        blocks_per_news = 10

    for index, text in updates.items():
        _replace_paragraph_text(doc.paragraphs[index], text)

    if len(doc.paragraphs) >= 169:
        _replace_cover_prepared_line(doc.paragraphs[31], data["tenant"]["name"], data["report"]["prepared_by"])

    for offset in range(action_slots):
        action_text = data["actions_worked"][offset] if offset < len(data["actions_worked"]) else ""
        _replace_paragraph_text(doc.paragraphs[action_start + offset], action_text)

    for news_idx in range(2):
        news = data["security_news"][news_idx] if news_idx < len(data["security_news"]) else None
        base = news_paragraph_start + (news_idx * blocks_per_news)
        if news is None:
            for paragraph_idx in range(base, min(base + blocks_per_news, len(doc.paragraphs))):
                _replace_paragraph_text(doc.paragraphs[paragraph_idx], "")
            continue
        _replace_paragraph_text(doc.paragraphs[base], news["title"])
        _replace_paragraph_text(doc.paragraphs[base + 1], f"Fuente: {news['source']}")
        _replace_paragraph_text(doc.paragraphs[base + 2], news["summary"])
        _replace_paragraph_text(doc.paragraphs[base + 3], f"Fecha de referencia: {news['date']}.")
        _replace_paragraph_text(doc.paragraphs[base + 4], f"Enlaces: {', '.join(news['links'])}")
        _replace_paragraph_text(doc.paragraphs[base + 5], "Se recomienda evaluar impacto, exposición local y controles compensatorios vigentes antes de definir la remediación.")
        _replace_paragraph_text(doc.paragraphs[base + 6], "El contenido se incluye como insumo contextual para el analista y no reemplaza validaciones internas ni gestión formal de cambios.")
        for paragraph_idx in range(base + 7, min(base + blocks_per_news, len(doc.paragraphs))):
            _replace_paragraph_text(doc.paragraphs[paragraph_idx], "")


def _build_data_base_paragraph(data: dict) -> str:
    return (
        f"{data['report']['prepared_by']} implementó este servicio integral de monitoreo y gestión de vulnerabilidades "
        f"con base en el análisis técnico del entorno evaluado, permitiendo consolidar la exposición por dominio y priorizar "
        f"los activos con mayor criticidad operativa."
    )


def _build_comparative_text(data: dict) -> str:
    sev = data["severity_summary"]
    return (
        f"En la semana actual se registran {sev.get('critical', 0)} hallazgos críticos, {sev.get('high', 0)} altos, "
        f"{sev.get('medium', 0)} medios, {sev.get('low', 0)} bajos y "
        f"{sev.get('informational', 0)} informativos. El comparativo semanal debe revisarse junto con la tendencia histórica "
        "para priorizar mitigaciones y validar si los controles implementados reducen la exposición operacional."
    )


def _build_histogram_text() -> str:
    return (
        "El histograma evidencia la distribución residual por severidad, concentrando mayor volumen en hallazgos informativos y medios, mientras que las "
        "severidades críticas y altas permanecen como prioridad para remediación, endurecimiento y validación operativa."
    )


def _build_results_summary() -> str:
    return (
        "El análisis de vulnerabilidades permitió consolidar la exposición por dominio, ordenar hallazgos por criticidad y mantener una vista semanal utilizable "
        "por el analista para seguimiento, priorización y coordinación de próximos pasos."
    )


def _build_requirement_text(data: dict) -> str:
    return (
        f"Se requiere revisar y priorizar los hallazgos críticos y altos consolidados en el informe {data['report']['id']}, "
        f"definiendo responsables, ventanas de intervención y criterios de cierre para el periodo reportado."
    )


def replace_headers_and_footers(doc: Document, data: dict) -> None:
    tenant_name = data["tenant"]["name"]
    prepared_by = data["report"]["prepared_by"]
    footer_text = f"{prepared_by}                                              Minority Report XOC\t{tenant_name}"

    for section in doc.sections:
        for paragraph in section.footer.paragraphs:
            if paragraph.text.strip():
                _replace_paragraph_text(paragraph, footer_text)


def populate_domain_tables(doc: Document, data: dict) -> None:
    findings_by_domain = _group_findings_by_domain(data)
    if len(doc.tables) >= 10:
        table_domain_pairs = [
            (1, "Dominio de Web Externo"),
            (2, "Dominio de IPs Públicas"),
            (3, "Dominio de FW"),
            (4, "Dominio Infraestructura de Computo - Servers"),
            (5, "Dominio Infraestructura de Red - Switches"),
            (6, "Dominio Infraestructura de Red - WIFI"),
            (7, "Dominio Infraestructura de Computo - Desktops"),
            (8, "Dominio Infraestructura de Computo - Desktops"),
            (9, "Dominio Infraestructura OT/IoT"),
        ]
    else:
        table_domain_pairs = list(enumerate(DOMAIN_TABLE_ORDER, start=1))

    populated_desktop_split = False
    for table_idx, domain_name in table_domain_pairs:
        table = doc.tables[table_idx]
        _clear_table_rows(table)
        domain_findings = findings_by_domain.get(_normalize_domain(domain_name), [])
        if len(doc.tables) >= 10 and domain_name == "Dominio Infraestructura de Computo - Desktops":
            if populated_desktop_split:
                continue
            populated_desktop_split = True
        if not domain_findings:
            _add_finding_row(
                table,
                {
                    "id": "N/A",
                    "title": "Sin hallazgos relevantes en este corte",
                    "affected_hosts": "0",
                    "severity": "Informativo",
                },
            )
            continue
        for finding in domain_findings:
            _add_finding_row(table, finding)


def generate_from_real_template(template_path: str, data: dict, output_path: str) -> str:
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    shutil.copy2(template_path, output_path)
    doc = Document(output_path)
    replace_known_paragraphs(doc, data)
    replace_headers_and_footers(doc, data)
    populate_domain_tables(doc, data)
    doc.save(output_path)
    return output_path


def _insert_table_after(paragraph, rows: int, cols: int):
    table = paragraph._parent.add_table(rows=rows, cols=cols, width=Inches(7.2))
    paragraph._p.addnext(table._tbl)
    return table


def insert_findings_table(rendered_path: str, data: dict, output_path: str) -> None:
    doc = Document(rendered_path)
    findings = data.get("findings", [])
    if not findings:
        doc.save(output_path)
        return
    columns = [
        ("ID", "id"),
        ("Dominio", "domain"),
        ("Vulnerabilidad", "title"),
        ("Hosts", "affected_hosts"),
        ("Severidad", "severity"),
    ]

    target_paragraph = None
    for paragraph in doc.paragraphs:
        if "[[FINDINGS_TABLE]]" in paragraph.text:
            target_paragraph = paragraph
            break

    if target_paragraph is None:
        doc.save(output_path)
        return

    table = _insert_table_after(target_paragraph, rows=1, cols=len(columns))
    table.style = "Table Grid"

    for col_idx, (column_name, _) in enumerate(columns):
        _set_cell_text(table.rows[0].cells[col_idx], column_name, bold=True)

    for record in findings:
        row_cells = table.add_row().cells
        for col_idx, (_, field_name) in enumerate(columns):
            _set_cell_text(row_cells[col_idx], str(record.get(field_name, "")))

    target_paragraph.text = ""
    doc.save(output_path)
    if rendered_path != output_path and os.path.exists(rendered_path):
        os.remove(rendered_path)


def validate_docx_file(docx_path: str) -> tuple[bool, str]:
    with zipfile.ZipFile(docx_path, "r") as archive:
        if "word/document.xml" not in archive.namelist():
            return False, "No existe word/document.xml en el DOCX generado"
        document_xml = archive.read("word/document.xml")
    etree.fromstring(document_xml)
    return True, "Validacion ZIP/XML OK"


def replace_residual_template_text(docx_path: str, data: dict) -> None:
    tenant_name = data["tenant"]["name"]
    xml_prefix = "word/"
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

    fd, tmp_name = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    tmp_path = tmp_name

    try:
        with zipfile.ZipFile(docx_path, "r") as source_zip, zipfile.ZipFile(tmp_path, "w", compression=zipfile.ZIP_DEFLATED) as target_zip:
            for item in source_zip.infolist():
                payload = source_zip.read(item.filename)
                if item.filename.startswith(xml_prefix) and item.filename.endswith(".xml"):
                    try:
                        root = etree.fromstring(payload)
                    except etree.XMLSyntaxError:
                        target_zip.writestr(item, payload)
                        continue

                    changed = False
                    for text_node in root.xpath(".//w:t", namespaces=ns):
                        text = text_node.text or ""
                        if text == "JOCKEY":
                            text_node.text = tenant_name
                            changed = True
                        elif text == "SALUD":
                            text_node.text = ""
                            changed = True
                        elif text in {"Jockey Salud", "JOCKEY SALUD", "Jockey Salud."}:
                            suffix = "." if text.endswith(".") else ""
                            text_node.text = f"{tenant_name}{suffix}"
                            changed = True

                    if changed:
                        payload = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")

                target_zip.writestr(item, payload)

        shutil.move(tmp_path, docx_path)
    finally:
        if os.path.exists(tmp_path):
            os.remove(tmp_path)


def generate_report_docx(
    template_path: str,
    data: dict,
    output_path: str,
    has_real_template: bool = False,
) -> str:
    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    if has_real_template:
        return generate_from_real_template(template_path, data, output_path)

    render_from_template(template_path, data, output_path)
    return output_path

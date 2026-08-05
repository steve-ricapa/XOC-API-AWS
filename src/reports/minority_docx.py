"""Genera Minority Report DOCX usando Azure AI Foundry + plantilla XOC.

La IA devuelve JSON; este script genera el DOCX final.
"""

from __future__ import annotations

import argparse
import json
import re
import unicodedata
import zipfile
from copy import deepcopy
from datetime import date
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table
from docx.text.paragraph import Paragraph
ROOT = Path(__file__).resolve().parent
REPOSITORY_ROOT = ROOT.parent.parent
TEMPLATES_DIR = REPOSITORY_ROOT / "Plantillas"
CLIENT_TEMPLATE_PATH = TEMPLATES_DIR / "TXDX_report_for_client-tenant MINORITY_REPORT_SEMANAL_dia_mes.docx"
ADMIN_CLIENT_TEMPLATE_PATH = TEMPLATES_DIR / "TXDX_report_for_admin_client-tenant MINORITY_REPORT_SEMANAL_dia_mes.docx"
TEMPLATE_PATH = CLIENT_TEMPLATE_PATH
OUTPUT_DIR = ROOT / "output"
BODY_START_MARKER = "[[BODY_START]]"
DOMAIN_TABLE_MARKER = "[[PROTO:DOMAIN_TABLE]]"
COMPARISON_TABLE_MARKER = "[[PROTO:COMPARISON_TABLE]]"
COVER_CLIENT_MARKER = "[[COVER_CLIENT]]"
COVER_PERIOD_MARKER = "[[COVER_PERIOD]]"
COVER_PREPARED_BY_MARKER = "[[COVER_PREPARED_BY]]"
FOOTER_CLIENT_MARKER = "[[FOOTER_CLIENT]]"

BRAND_GREEN = RGBColor(0x00, 0xFF, 0x9F)
BRAND_BLUE = RGBColor(0x00, 0xF0, 0xFF)
MODEL_BLUE = RGBColor(0x00, 0x6D, 0x9F)
DARK_TEXT = RGBColor(0x1F, 0x38, 0x62)
MUTED_TEXT = RGBColor(0x5C, 0x66, 0x70)
SOC_FILL = "FFE5E5"
NOC_FILL = "E6F7F5"
SOC_HEADER_FILL = "8B1E3F"
NOC_HEADER_FILL = "0F766E"
SOC_ALT_FILL = "FFF5F6"
NOC_ALT_FILL = "F2FBFA"
SOC_SEVERITY_ROW_FILLS = {
    "ALTO": "FDE7EA",
    "MEDIO": "FFF4DA",
    "BAJO": "F4FBF7",
}
NOC_SEVERITY_ROW_FILLS = {
    "ALTO": "E6F7F5",
    "MEDIO": "EEF7FF",
    "BAJO": "F5FAF9",
}
TABLE_HEADER = "0B1F2A"
TABLE_ALT = "EAF8F5"
CALLOUT = "EAF8F5"
WARNING = "FFF4D6"
BORDER = "8FDACC"
SEVERITY_COLOR_MAP = {
    "CRÍTICA": "B00020",
    "CRITICA": "B00020",
    "ALTA": "E53935",
    "ALTO": "E53935",
    "MEDIA": "FBC02D",
    "MEDIO": "FBC02D",
    "BAJA": "A5D6A7",
    "BAJO": "A5D6A7",
    "INFORMATIVA": "43A047",
    "INFORMATIVO": "43A047",
}


SECTION_ANCHORS = {
    "general": ["1. Datos generales", "Datos generales"],
    "service": ["1.1. Servicio de Monitoreo", "Servicio de Monitoreo"],
    "period": ["1.2. Periodo", "Periodo"],
    "tools": ["1.3. Herramientas", "Herramientas"],
    "data_base": ["1.4. Datos Base", "Datos Base"],
    "executive": ["2. Resumen ejecutivo del dominio", "Resumen ejecutivo del dominio"],
    "comparison": [
        "2.1. Distribución actual de hallazgos por severidad",
        "Distribución actual de hallazgos por severidad",
    ],
    "histogram": ["2.2. Estado actual de la seguridad", "Estado actual de la seguridad"],
    "results_next": ["2.3. Resultados obtenidos y próximas acciones", "Resultados obtenidos y próximas acciones"],
    "results_obtained": ["2.4. Resultados obtenidos", "Resultados obtenidos"],
    "next_actions": ["2.5. Próximas acciones", "Próximas acciones"],
    "requirements": ["2.5.1. Requerimiento", "Requerimiento"],
    "domains": ["3. Seguridad por Dominio", "Seguridad por Dominio"],
    "weekly_actions": [
        "4. Reporte de acciones trabajadas durante la semana",
        "Reporte de acciones trabajadas durante la semana",
    ],
    "results": ["5. Resultados obtenidos", "Resultados obtenidos"],
    "reinforced_security": ["5.1. Seguridad Reforzada", "Seguridad Reforzada"],
    "pending_findings": ["5.2. Hallazgos pendientes", "Hallazgos pendientes"],
    "security_news": ["6. Noticias de seguridad", "Noticias de seguridad"],
}


SECTION_ORDER = [
    "general",
    "service",
    "period",
    "tools",
    "data_base",
    "executive",
    "comparison",
    "histogram",
    "results_next",
    "results_obtained",
    "next_actions",
    "requirements",
    "domains",
    "weekly_actions",
    "results",
    "reinforced_security",
    "pending_findings",
    "security_news",
]


def normalize_report_variant(value: str | None) -> str:
    normalized = (value or "client").strip().lower().replace("-", "_")
    if normalized in {"admin", "client_admin", "admin_client", "report_for_admin_client"}:
        return "client_admin"
    return "client"


def base_template_path(report_variant: str | None = None) -> Path:
    variant = normalize_report_variant(report_variant)
    candidates = (
        [ADMIN_CLIENT_TEMPLATE_PATH, TEMPLATE_PATH] if variant == "client_admin" else [CLIENT_TEMPLATE_PATH, TEMPLATE_PATH]
    )
    for path in candidates:
        if path.is_file():
            return path
    raise FileNotFoundError(f"No existe plantilla base para Minority Report ({variant}).")


def _safe_name(value: str) -> str:
    value = re.sub(r"[^A-Za-z0-9_-]+", "-", value.upper().strip())
    return value.strip("-")[:80] or "CLIENTE"


def _output_path(payload: dict[str, Any]) -> Path:
    document_code = _clean_text(payload.get("document_code"))
    if document_code:
        return OUTPUT_DIR / f"{_safe_name(document_code)}.docx"
    client = _safe_name(payload.get("client_name") or "CLIENTE")
    period = _safe_name(payload.get("period") or date.today().isoformat())
    return OUTPUT_DIR / f"MINORITY-REPORT-XOC_{client}_{period}.docx"


def build_output_filename(payload: dict[str, Any]) -> str:
    document_code = _clean_text(payload.get("document_code"))
    if document_code:
        return f"{_safe_name(document_code)}.docx"
    client = _safe_name(payload.get("client_name") or "CLIENTE")
    period = _safe_name(payload.get("period") or date.today().isoformat())
    return f"MINORITY_REPORT_XOC_{client}_{period}.docx"


def _writable_output_path(path: Path) -> Path:
    if not path.exists():
        return path
    try:
        with path.open("r+b"):
            return path
    except PermissionError:
        for version in range(2, 100):
            candidate = path.with_name(f"{path.stem}_v{version}{path.suffix}")
            if not candidate.exists():
                return candidate
        raise


def _as_list(value: Any) -> list[Any]:
    if isinstance(value, list):
        return value
    return [value] if value not in (None, "") else []


def _clean_text(value: Any) -> str:
    return str(value or "").strip()


def _textbox_text(box: Any) -> str:
    return "".join(node.text or "" for node in box.iter(qn("w:t")))


def _set_textbox_text(box: Any, value: str, *, size_pt: float | None = None) -> None:
    nodes = list(box.iter(qn("w:t")))
    if not nodes:
        return
    nodes[0].text = value
    for node in nodes[1:]:
        node.text = ""
    if size_pt:
        half_points = str(int(size_pt * 2))
        for run in box.iter(qn("w:r")):
            r_pr = run.find(qn("w:rPr"))
            if r_pr is None:
                r_pr = OxmlElement("w:rPr")
                run.insert(0, r_pr)
            size = r_pr.find(qn("w:sz"))
            if size is None:
                size = OxmlElement("w:sz")
                r_pr.append(size)
            size.set(qn("w:val"), half_points)
            size_cs = r_pr.find(qn("w:szCs"))
            if size_cs is None:
                size_cs = OxmlElement("w:szCs")
                r_pr.append(size_cs)
            size_cs.set(qn("w:val"), half_points)


def _replace_textbox_jockey_client(box: Any, client: str) -> None:
    nodes = list(box.iter(qn("w:t")))
    for index, node in enumerate(nodes):
        if (node.text or "").strip() == "JOCKEY":
            node.text = client
            # El modelo tiene "JOCKEY" + espacio + "SALUD"; se limpian esos
            # fragmentos para no dejar restos del cliente ejemplo.
            for cleanup in nodes[index + 1 : index + 4]:
                if (cleanup.text or "").strip() in {"", "SALUD"}:
                    cleanup.text = ""
            return


def _clear_textbox_content(box: Any) -> None:
    for node in box.iter(qn("w:t")):
        node.text = ""


def _set_cover_line(paragraph: Paragraph, text: str) -> None:
    paragraph.text = ""
    run = paragraph.add_run(text)
    run.font.name = "Lucida Sans Unicode"
    run.font.size = Pt(14)
    run.font.color.rgb = DARK_TEXT


def _clear_story(story: Any) -> None:
    root = story._element
    for child in list(root):
        root.remove(child)


def _append_field_run(paragraph: Paragraph, instruction: str) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    paragraph.add_run()._r.append(begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {instruction} "
    paragraph.add_run()._r.append(instr)

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    paragraph.add_run()._r.append(separate)

    paragraph.add_run("1")

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    paragraph.add_run()._r.append(end)


def _set_header_page_counter(story: Any) -> None:
    _clear_story(story)
    paragraph = story.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    prefix = paragraph.add_run("Página ")
    prefix.font.size = Pt(9)
    prefix.font.color.rgb = MUTED_TEXT
    _append_field_run(paragraph, "PAGE")
    middle = paragraph.add_run(" de ")
    middle.font.size = Pt(9)
    middle.font.color.rgb = MUTED_TEXT
    _append_field_run(paragraph, "NUMPAGES")
    for run in paragraph.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = MUTED_TEXT


def update_cover_and_footer(document: Document, payload: dict[str, Any]) -> None:
    client = _clean_text(payload.get("client_name")) or "Cliente"
    period = _clean_text(payload.get("period")) or "Periodo no especificado"
    prepared_by = _clean_text(payload.get("prepared_by")) or "TXDXSECURE"

    for box in document._element.xpath(".//w:txbxContent"):
        original = _textbox_text(box).strip()
        if original in {"Change for date", COVER_PERIOD_MARKER}:
            _set_textbox_text(box, period, size_pt=9.5)
        elif original in {"Change for prepared for", COVER_CLIENT_MARKER}:
            _set_textbox_text(box, client, size_pt=11)
        elif original in {"TXDXSECURE", COVER_PREPARED_BY_MARKER}:
            _set_textbox_text(box, prepared_by, size_pt=11)
        elif "JOCKEY" in original and "SALUD" in original:
            _replace_textbox_jockey_client(box, client)
        elif original == FOOTER_CLIENT_MARKER:
            _set_textbox_text(box, client, size_pt=11)

    for paragraph in document.paragraphs[:45]:
        text = paragraph.text.strip()
        if "JOCKEY SALUD" in text and "TXDXSECURE" in text:
            _set_cover_line(paragraph, f"{client}\t{prepared_by}")
        elif COVER_CLIENT_MARKER in text or COVER_PREPARED_BY_MARKER in text:
            _set_cover_line(paragraph, text.replace(COVER_CLIENT_MARKER, client).replace(COVER_PREPARED_BY_MARKER, prepared_by))
        elif "Del 20 de junio al 26 de junio del 2026" in text or "Del 20 al 26 de junio del 2026" in text:
            _set_cover_line(paragraph, period)
        elif COVER_PERIOD_MARKER in text:
            _set_cover_line(paragraph, period)

    for section_index, section in enumerate(document.sections):
        if not (section_index == 0 and len(document.sections) > 1):
            _set_header_page_counter(section.header)
            if hasattr(section, "first_page_header"):
                _set_header_page_counter(section.first_page_header)
            if hasattr(section, "even_page_header"):
                _set_header_page_counter(section.even_page_header)
        footer = section.footer
        footer.is_linked_to_previous = False
        if not footer.paragraphs:
            footer.add_paragraph()
        paragraph = footer.paragraphs[0]
        paragraph.text = ""
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.tab_stops.clear_all()
        paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(2.9))
        paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(6.15))
        for index, text in enumerate(("TxdxSecure", "Minority Report XOC", client)):
            if index:
                paragraph.add_run("\t")
            run = paragraph.add_run(text)
            run.font.size = Pt(8.5)
            run.font.color.rgb = MUTED_TEXT

        # Limpia párrafos adicionales heredados del template que dejan placeholders visibles.
        for extra in footer.paragraphs[1:]:
            extra.text = ""


def clear_template_body_after_cover(document: Document) -> None:
    """Conserva portada/TOC/marker y remueve cualquier prototipo o cuerpo restante."""
    body = document._element.body
    children = list(body)
    preserve_until = -1
    for index, child in enumerate(children):
        text = "".join(node.text or "" for node in child.iter(qn("w:t"))).strip()
        if BODY_START_MARKER in text:
            preserve_until = index
            break
        if text == "Datos generales" and preserve_until < 0:
            preserve_until = index - 1
            break
    if preserve_until >= 0:
        while preserve_until + 1 < len(children):
            next_child = children[preserve_until + 1]
            next_text = "".join(node.text or "" for node in next_child.iter(qn("w:t"))).strip()
            has_section_break = bool(next_child.findall(".//" + qn("w:sectPr")))
            if next_text:
                break
            if has_section_break:
                preserve_until += 1
                break
            preserve_until += 1
    if preserve_until < 0:
        last_drawing_index = -1
        for index, child in enumerate(children):
            if child.findall(".//" + qn("w:drawing")):
                last_drawing_index = index
        preserve_until = last_drawing_index if last_drawing_index >= 0 else -1
    preserved = set(children[: preserve_until + 1])
    for child in list(body):
        if child not in preserved and child.tag != "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr":
            body.remove(child)


def _style_name(document: Document, *names: str) -> str:
    for name in names:
        if name in document.styles:
            return name
    return names[-1]


def _format_run(
    run: Any,
    *,
    bold: bool = False,
    italic: bool = False,
    color: RGBColor | None = None,
    size: float | None = None,
    font_name: str | None = None,
) -> None:
    run.bold = bold or None
    run.italic = italic or None
    if font_name:
        run.font.name = font_name
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)


def apply_example_body_style(document: Document) -> None:
    """Prepara estilos auxiliares sin modificar estilos globales de la portada."""
    styles = document.styles
    if "Body Text" not in styles:
        styles.add_style("Body Text", WD_STYLE_TYPE.PARAGRAPH)

def apply_body_page_layout(document: Document) -> None:
    """Centra el cuerpo para impresión sin tocar la primera sección/portada."""
    body_sections = list(document.sections)[1:] or []
    for section in body_sections:
        section.left_margin = Inches(1.05)
        section.right_margin = Inches(1.05)
        section.top_margin = Inches(0.86)
        section.bottom_margin = Inches(0.58)


def _spacing(paragraph: Any, *, before: float = 0, after: float = 7, left: float = 0, first: float | None = None) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.left_indent = Pt(left) if left else None
    if first is not None:
        paragraph.paragraph_format.first_line_indent = Pt(first)


def _set_paragraph_bottom_border(paragraph: Any, color: str = "00AEEF", size: str = "8") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)


def add_heading(document: Document, title: str, number: str | None = None) -> None:
    paragraph = document.add_paragraph()
    paragraph.style = _style_name(document, "Heading 2", "Normal")
    _spacing(paragraph, before=12, after=6)
    display = f"{number} {title}" if number else title
    paragraph.add_run(display)


def add_subheading(document: Document, title: str, number: str | None = None) -> None:
    paragraph = document.add_paragraph()
    paragraph.style = _style_name(document, "Heading 3", "Normal")
    _spacing(paragraph, before=7, after=4)
    display = f"{number} {title}" if number else title
    paragraph.add_run(display)


def add_body(document: Document, text: Any, *, indent: bool = False) -> None:
    if isinstance(text, list):
        chunks = [_clean_text(item) for item in text if _clean_text(item)]
    else:
        chunks = [chunk.strip() for chunk in str(text or "").splitlines() if chunk.strip()]
    for chunk in chunks:
        paragraph = document.add_paragraph()
        paragraph.style = _style_name(document, "Body Text", "Normal")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _spacing(paragraph, after=5, left=18 if indent else 0)
        paragraph.add_run(chunk)


def add_bullets(document: Document, values: Any) -> None:
    for value in _as_list(values):
        text = _clean_text(value)
        if not text:
            continue
        paragraph = document.add_paragraph()
        paragraph.style = _style_name(document, "List Paragraph", "Body Text", "Normal")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _spacing(paragraph, after=4)
        paragraph.add_run(text)


def _shade_cell(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_border(cell: Any, color: str = BORDER, size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def _set_cell_width(cell: Any, width_inches: float) -> None:
    cell.width = Inches(width_inches)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))


def _set_row_widths(row: Any, widths: list[float]) -> None:
    for cell, width in zip(row.cells, widths):
        _set_cell_width(cell, width)


def _set_cell_text(cell: Any, text: Any, *, bold: bool = False, color: RGBColor | None = None, size: float = 9.2) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    _spacing(paragraph, after=2)
    run = paragraph.add_run(str(text or ""))
    _format_run(run, bold=bold, color=color, size=size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _shade_row(row: Any, fill: str) -> None:
    for cell in row.cells:
        _shade_cell(cell, fill)


def _domain_table_palette(layer: str) -> dict[str, Any]:
    normalized = (layer or "SOC").upper()
    if normalized == "NOC":
        return {
            "header_fill": NOC_HEADER_FILL,
            "row_alt_fill": NOC_ALT_FILL,
            "severity_fills": NOC_SEVERITY_ROW_FILLS,
        }
    return {
        "header_fill": SOC_HEADER_FILL,
        "row_alt_fill": SOC_ALT_FILL,
        "severity_fills": SOC_SEVERITY_ROW_FILLS,
    }


def add_callout(document: Document, title: str, content: Any = "", *, fill: str = CALLOUT, bullets: Any = None) -> None:
    text = _clean_text(content)
    bullet_values = [_clean_text(item) for item in _as_list(bullets) if _clean_text(item)] if bullets is not None else []
    if not text and not bullet_values:
        return
    add_subheading(document, title)
    if text:
        add_body(document, text)
    if bullet_values:
        add_bullets(document, bullet_values)


def add_key_value_table(document: Document, rows: list[tuple[str, Any]]) -> None:
    table = document.add_table(rows=1, cols=2)
    table.autofit = False
    widths = [1.25, 4.25]
    try:
        table.style = "Table Normal"
    except KeyError:
        pass
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_row_widths(table.rows[0], widths)
    for cell, header in zip(table.rows[0].cells, ("Campo", "Detalle")):
        _shade_cell(cell, "006D9F")
        _set_cell_border(cell, "BFBFBF", "4")
        _set_cell_text(cell, header, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    for index, (key, value) in enumerate(rows, start=1):
        cells = table.add_row().cells
        _set_row_widths(table.rows[-1], widths)
        for cell in cells:
            _set_cell_border(cell, "BFBFBF", "4")
        _set_cell_text(cells[0], key, bold=True, color=DARK_TEXT)
        _set_cell_text(cells[1], value)
    document.add_paragraph()


def add_findings_table(document: Document, findings: list[dict[str, Any]], *, layer: str = "SOC") -> None:
    if not findings:
        return
    palette = _domain_table_palette(layer)
    table = document.add_table(rows=1, cols=4)
    table.autofit = False
    widths = [1.02, 2.82, 0.92, 0.72]
    try:
        table.style = "Table Normal"
    except KeyError:
        pass
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_row_widths(table.rows[0], widths)
    for cell, header in zip(table.rows[0].cells, ("ID", "Vulnerabilidades", "Host Afectados", "Severidad")):
        _shade_cell(cell, palette["header_fill"])
        _set_cell_border(cell, "BFBFBF", "4")
        _set_cell_text(cell, header, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=8.8)
    for index, finding in enumerate(findings, start=1):
        cells = table.add_row().cells
        _set_row_widths(table.rows[-1], widths)
        for cell in cells:
            _set_cell_border(cell, "BFBFBF", "4")
        severity = _display_finding_severity(finding.get("severity"))
        fill = palette["row_alt_fill"] if index % 2 == 0 else "FFFFFF"
        fill = palette["severity_fills"].get(severity, fill)
        _shade_row(table.rows[-1], fill)
        _set_cell_text(cells[0], finding.get("id"), bold=False, color=DARK_TEXT, size=8.5)
        _set_cell_text(cells[1], finding.get("vulnerability"), size=8.5)
        _set_cell_text(cells[2], finding.get("affected_hosts"), size=8.5)
        severity_color = RGBColor(0xC0, 0x00, 0x00) if severity == "ALTO" else DARK_TEXT
        _set_cell_text(cells[3], severity, bold=True, color=severity_color, size=8.5)
    document.add_paragraph()


def _remove_table_data_rows(table: Table) -> None:
    """Conserva el encabezado corporativo de la plantilla y elimina su data demo."""
    for row in list(table.rows[1:]):
        table._tbl.remove(row._tr)


def _populate_template_findings_table(table: Table, findings: list[dict[str, Any]], *, layer: str = "SOC") -> None:
    """Llena una tabla heredada de la plantilla sin recrear su formato OOXML."""
    if not table.rows or len(table.columns) != 4:
        raise ValueError("La tabla de hallazgos de la plantilla debe tener cuatro columnas")

    palette = _domain_table_palette(layer)
    headers = ("ID", "Vulnerabilidades", "Host Afectados", "Severidad")
    for cell, header in zip(table.rows[0].cells, headers):
        _shade_cell(cell, palette["header_fill"])
        _set_cell_text(cell, header, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=8.8)

    _remove_table_data_rows(table)
    rows = findings or []
    for index, finding in enumerate(rows):
        cells = table.add_row().cells
        fill = palette["row_alt_fill"] if index % 2 else "FFFFFF"
        severity = _display_finding_severity(finding.get("severity"))
        fill = palette["severity_fills"].get(severity, fill)
        _shade_row(table.rows[-1], fill)
        _set_cell_text(cells[0], finding.get("id"), color=DARK_TEXT, size=8.5)
        _set_cell_text(cells[1], finding.get("vulnerability"), size=8.5)
        _set_cell_text(cells[2], finding.get("affected_hosts"), size=8.5)
        severity_color = RGBColor(0xC0, 0x00, 0x00) if severity == "ALTO" else DARK_TEXT
        _set_cell_text(cells[3], severity, bold=True, color=severity_color, size=8.5)


def _populate_template_comparison_table(table: Table, rows: list[dict[str, Any]]) -> None:
    if not table.rows or len(table.columns) != 3:
        raise ValueError("La tabla comparativa de la plantilla debe tener tres columnas")
    headers = ("Severidad", "Referencia anterior", "Estado actual")
    for cell, header in zip(table.rows[0].cells, headers):
        _set_cell_text(cell, header, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=8.8)
    _remove_table_data_rows(table)
    for row in rows:
        cells = table.add_row().cells
        severity = row.get("severity")
        _shade_cell(cells[0], _severity_color_fill(severity))
        _set_cell_text(cells[0], severity, bold=True, color=_severity_text_color(severity), size=8.5)
        _set_cell_text(cells[1], row.get("previous"), size=8.5)
        _set_cell_text(cells[2], row.get("current"), size=8.5)


def capture_domain_table_templates(document: Document) -> list[Any]:
    """Captura las tablas del modelo antes de limpiar su contenido de ejemplo.

    Se reinsertan como clones, de modo que se mantiene el diseño corporativo
    (bordes, rellenos y anchos) en lugar de generar tablas desde cero.
    """
    marker_index = next((index for index, p in enumerate(document.paragraphs) if DOMAIN_TABLE_MARKER in p.text), None)
    if marker_index is not None:
        for table in document.tables:
            if len(table.columns) == 4 and table.rows:
                return [deepcopy(table._tbl)]
    return [deepcopy(table._tbl) for table in document.tables if len(table.columns) == 4 and table.rows]


def capture_comparison_table_template(document: Document) -> Any | None:
    marker_index = next((index for index, p in enumerate(document.paragraphs) if COMPARISON_TABLE_MARKER in p.text), None)
    if marker_index is not None:
        for table in document.tables:
            if len(table.columns) == 3 and table.rows:
                return deepcopy(table._tbl)
    for table in document.tables:
        if len(table.columns) == 3 and table.rows:
            return deepcopy(table._tbl)
    return None


def add_template_findings_table(document: Document, template_table: Any, findings: list[dict[str, Any]], *, layer: str = "SOC") -> None:
    marker = document.add_paragraph()
    table = Table(deepcopy(template_table), document)
    marker._p.addprevious(table._tbl)
    _populate_template_findings_table(table, findings, layer=layer)
    document.add_paragraph()


def add_coverage_table(document: Document, rows: list[dict[str, Any]]) -> None:
    if not rows:
        return
    table = document.add_table(rows=1, cols=5)
    table.autofit = False
    widths = [1.9, 0.9, 1.65, 0.95, 1.35]
    try:
        table.style = "Table Normal"
    except KeyError:
        pass
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_row_widths(table.rows[0], widths)
    headers = ("Integración", "Capa", "Última evidencia", "Hallazgos", "Estado")
    for cell, header in zip(table.rows[0].cells, headers):
        _shade_cell(cell, "006D9F")
        _set_cell_border(cell, "BFBFBF", "4")
        _set_cell_text(cell, header, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=8.6)
    for row in rows:
        cells = table.add_row().cells
        _set_row_widths(table.rows[-1], widths)
        for cell in cells:
            _set_cell_border(cell, "BFBFBF", "4")
        _set_cell_text(cells[0], row.get("integration"), size=8.4)
        _set_cell_text(cells[1], row.get("layer"), bold=True, color=DARK_TEXT, size=8.4)
        _set_cell_text(cells[2], row.get("last_evidence_at"), size=8.4)
        _set_cell_text(cells[3], row.get("current_findings_total"), size=8.4)
        _set_cell_text(cells[4], row.get("status"), size=8.4)
    document.add_paragraph()


def add_domain_snapshot_table(document: Document, rows: list[tuple[str, Any]], *, fill: str) -> None:
    table = document.add_table(rows=1, cols=len(rows))
    table.autofit = False
    widths = [max(1.0, 6.9 / max(1, len(rows)))] * len(rows)
    try:
        table.style = "Table Normal"
    except KeyError:
        pass
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_row_widths(table.rows[0], widths)
    for index, (label, value) in enumerate(rows):
        cell = table.rows[0].cells[index]
        _shade_cell(cell, fill)
        _set_cell_border(cell, "BFBFBF", "4")
        cell.text = ""
        title = cell.paragraphs[0]
        _spacing(title, after=1)
        title.add_run(str(label or ""))
        title.runs[0].bold = True
        detail = cell.add_paragraph()
        _spacing(detail, after=0)
        detail.add_run(str(value or "N/D"))
    document.add_paragraph()


def _domain_layer(domain: dict[str, Any]) -> str:
    layer = _clean_text(domain.get("layer")).upper()
    if layer in {"SOC", "NOC"}:
        return layer
    provider = _clean_text(domain.get("provider")).lower()
    return "NOC" if provider in {"zabbix", "uptime_kuma"} else "SOC"


def _domain_latest_snapshot(domain: dict[str, Any]) -> str:
    snapshot = domain.get("snapshot") or {}
    scanned_at = _clean_text(snapshot.get("scanned_at"))
    if scanned_at:
        return scanned_at[:10]
    return "No disponible"


def _domain_soc_kpis(domain: dict[str, Any]) -> list[tuple[str, Any]]:
    severity = domain.get("current_severity_summary") or {}
    return [
        ("Capa", "SOC"),
        ("Último snapshot", _domain_latest_snapshot(domain)),
        ("Hallazgos", domain.get("current_findings_total") or 0),
        ("Críticos + Altos", int(severity.get("critical", 0) or 0) + int(severity.get("high", 0) or 0)),
    ]


def _domain_noc_kpis(domain: dict[str, Any]) -> list[tuple[str, Any]]:
    severity = domain.get("current_severity_summary") or {}
    status = "Con eventos observables" if int(domain.get("current_findings_total") or 0) > 0 else "Cobertura sin eventos"
    return [
        ("Capa", "NOC"),
        ("Último snapshot", _domain_latest_snapshot(domain)),
        ("Eventos", domain.get("current_findings_total") or 0),
        ("Estado", status),
        ("Altos + Medios", int(severity.get("high", 0) or 0) + int(severity.get("medium", 0) or 0)),
    ]


def _domain_callout_title(domain: dict[str, Any]) -> str:
    return "Riesgo prioritario" if _domain_layer(domain) == "SOC" else "Impacto operativo"


def _domain_callout_text(domain: dict[str, Any]) -> str:
    severity = domain.get("current_severity_summary") or {}
    if _domain_layer(domain) == "SOC":
        return (
            f"El dominio concentra {int(severity.get('critical', 0) or 0)} hallazgos críticos y "
            f"{int(severity.get('high', 0) or 0)} hallazgos altos, por lo que debe priorizarse en remediación y seguimiento técnico."
        )
    return (
        f"La última evidencia de este dominio refleja {domain.get('current_findings_total') or 0} eventos observables. "
        "La lectura debe enfocarse en continuidad, estabilidad e impacto sobre la operación monitoreada."
    )


def add_template_comparison_table(document: Document, template_table: Any, rows: list[dict[str, Any]]) -> None:
    marker = document.add_paragraph()
    table = Table(deepcopy(template_table), document)
    marker._p.addprevious(table._tbl)
    _populate_template_comparison_table(table, rows)
    document.add_paragraph()


def _display_finding_severity(value: Any) -> str:
    text = _clean_text(value).lower()
    if text in {"crítica", "critica", "alto", "alta", "high", "critical"}:
        return "ALTO"
    if text in {"media", "medio", "medium"}:
        return "MEDIO"
    return "BAJO"


def _severity_color_fill(value: Any) -> str:
    normalized = _plain_match_text(_clean_text(value)).upper()
    return SEVERITY_COLOR_MAP.get(normalized, "FFFFFF")


def _severity_text_color(value: Any) -> RGBColor:
    normalized = _plain_match_text(_clean_text(value)).upper()
    if normalized in {"MEDIA", "MEDIO", "BAJA", "BAJO"}:
        return RGBColor(0x1F, 0x38, 0x62)
    return RGBColor(0xFF, 0xFF, 0xFF)


def add_severity_comparison_table(document: Document, rows: list[dict[str, Any]]) -> None:
    if not rows:
        return
    table = document.add_table(rows=1, cols=3)
    table.autofit = False
    widths = [1.35, 1.25, 1.25]
    try:
        table.style = "Table Normal"
    except KeyError:
        pass
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_row_widths(table.rows[0], widths)
    for cell, header in zip(table.rows[0].cells, ("Severidad", "Referencia anterior", "Estado actual")):
        _shade_cell(cell, "006D9F")
        _set_cell_border(cell, "BFBFBF", "4")
        _set_cell_text(cell, header, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=8.8)
    for row in rows:
        cells = table.add_row().cells
        _set_row_widths(table.rows[-1], widths)
        for cell in cells:
            _set_cell_border(cell, "BFBFBF", "4")
        severity = row.get("severity")
        _shade_cell(cells[0], _severity_color_fill(severity))
        _set_cell_text(cells[0], severity, bold=True, color=_severity_text_color(severity), size=8.5)
        _set_cell_text(cells[1], row.get("previous"), size=8.5)
        _set_cell_text(cells[2], row.get("current"), size=8.5)
    document.add_paragraph()


def add_caption(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _spacing(paragraph, before=2, after=8)
    run = paragraph.add_run(text)
    _format_run(run, italic=True, color=MUTED_TEXT, size=8.5)


def enable_update_fields_on_open(document: Document) -> None:
    settings = document.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def add_word_toc_field(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.style = _style_name(document, "Normal")

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    paragraph.add_run()._r.append(begin)

    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = ' TOC \\o "1-3" \\h \\z \\u '
    paragraph.add_run()._r.append(instruction)

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    paragraph.add_run()._r.append(separate)

    placeholder = paragraph.add_run(
        "Índice automático. Al abrir el documento en Word, actualice los campos si no se muestran las páginas."
    )
    placeholder.italic = True
    placeholder.font.size = Pt(9)
    placeholder.font.color.rgb = MUTED_TEXT

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    paragraph.add_run()._r.append(end)


def _index_entries(report_variant: str, payload: dict[str, Any]) -> list[tuple[str, int]]:
    entries: list[tuple[str, int]] = [
        ("1. Datos generales", 0),
        ("1.1. Servicio de Monitoreo", 1),
        ("1.2. Periodo", 1),
        ("1.3. Herramientas", 1),
        ("1.4. Datos Base", 1),
        ("1.5. Cobertura del servicio", 1),
        ("2. Resumen ejecutivo del dominio", 0),
        ("2.1. Distribución actual de hallazgos por severidad", 1),
        ("2.2. Estado actual de la seguridad", 1),
        ("2.3. Focos prioritarios", 1),
        ("2.4. Consideraciones operativas", 1),
    ]
    if report_variant == "client_admin":
        entries.extend(
            [
                ("2.5. Resultados obtenidos y próximas acciones", 1),
                ("2.6. Resultados obtenidos", 1),
                ("2.7. Próximas acciones", 1),
                ("2.7.1. Requerimiento", 2),
            ]
        )
    entries.append(("3. Seguridad por Dominio", 0))
    for index, domain in enumerate(payload.get("security_domains") or [], start=1):
        name = _clean_text(domain.get("name")) if isinstance(domain, dict) else ""
        if name:
            entries.append((f"3.{index}. {name}", 1))
    if report_variant == "client_admin":
        entries.extend(
            [
                ("4. Reporte de acciones trabajadas durante la semana", 0),
                ("5. Resultados obtenidos", 0),
                ("5.1. Seguridad Reforzada", 1),
                ("5.2. Hallazgos pendientes", 1),
                ("6. Noticias de seguridad", 0),
            ]
        )
    return entries


def refresh_template_index(document: Document, payload: dict[str, Any], report_variant: str) -> None:
    """Reemplaza el índice dummy de la plantilla por uno que calce con el reporte generado."""
    body = document._element.body
    children = list(body)
    content_child = None
    for child in children:
        text = "".join(node.text or "" for node in child.iter(qn("w:t"))).strip()
        if text == "Contenido":
            content_child = child
            break
    if content_child is None:
        return

    current = content_child.getnext()
    index_sdt = current if current is not None and current.tag == qn("w:sdt") else None
    if index_sdt is None:
        return

    sdt_content = index_sdt.find(qn("w:sdtContent"))
    if sdt_content is None:
        return

    template_by_level: dict[int, Any] = {}
    for paragraph in sdt_content.findall(qn("w:p")):
        style = paragraph.find(".//" + qn("w:pStyle"))
        style_value = style.get(qn("w:val")) if style is not None else ""
        if style_value == "TDC2" and 0 not in template_by_level:
            template_by_level[0] = deepcopy(paragraph)
        elif style_value == "TDC1" and 1 not in template_by_level:
            template_by_level[1] = deepcopy(paragraph)
        elif style_value == "TDC3" and 2 not in template_by_level:
            template_by_level[2] = deepcopy(paragraph)

    first_template = sdt_content.find(qn("w:p"))
    fallback_template = deepcopy(first_template) if first_template is not None else OxmlElement("w:p")
    for child in list(sdt_content):
        sdt_content.remove(child)

    for text, level in _index_entries(report_variant, payload):
        source_template = template_by_level.get(level)
        if source_template is None:
            source_template = template_by_level.get(1)
        if source_template is None:
            source_template = fallback_template
        paragraph = deepcopy(source_template)
        text_nodes = list(paragraph.iter(qn("w:t")))
        if not text_nodes:
            run = paragraph.find(qn("w:r"))
            if run is None:
                run = OxmlElement("w:r")
                paragraph.append(run)
            text_node = OxmlElement("w:t")
            run.append(text_node)
            text_nodes = [text_node]
        text_nodes[0].text = text
        for node in text_nodes[1:]:
            node.text = ""
        sdt_content.append(paragraph)


def document_already_has_page_break(document: Document) -> bool:
    for paragraph in document.paragraphs:
        for br in paragraph._p.findall(".//" + qn("w:br")):
            if br.get(qn("w:type")) == "page":
                return True
    return False


def ensure_body_starts_after_cover(document: Document) -> None:
    """Evita la hoja de índice vacía y conserva el salto natural de la portada."""
    if len(document.sections) < 2 and not document_already_has_page_break(document):
        document.add_page_break()
    apply_body_page_layout(document)


def add_content_overview(document: Document, payload: dict[str, Any]) -> None:
    if not document_already_has_page_break(document):
        document.add_page_break()
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.style = _style_name(document, "Heading 1", "Normal")
    _spacing(title, before=8, after=10)
    run = title.add_run("Contenido")
    _format_run(run, bold=False)

    add_word_toc_field(document)
    document.add_page_break()


def build_report_body(
    document: Document,
    payload: dict[str, Any],
    domain_table_templates: list[Any] | None = None,
    comparison_table_template: Any | None = None,
) -> None:
    enable_update_fields_on_open(document)
    ensure_body_starts_after_cover(document)
    report_variant = normalize_report_variant(
        payload.get("report_variant")
        or payload.get("template_variant")
        or (payload.get("mock_meta") or {}).get("report_variant")
    )
    include_admin_sections = report_variant == "client_admin"
    refresh_template_index(document, payload, report_variant)
    for paragraph in document.paragraphs:
        if BODY_START_MARKER in paragraph.text:
            paragraph.text = ""
            break

    add_heading(document, "Datos generales", "1.")
    add_subheading(document, "Servicio de Monitoreo", "1.1.")
    service_text = (
        f"{payload.get('service_name') or 'Servicio de monitoreo proactivo XOC'} implementado por "
        f"{payload.get('prepared_by') or 'TXDXSECURE'} para el cliente {payload.get('client_name') or 'Cliente'}."
    )
    add_body(document, service_text)
    add_subheading(document, "Periodo", "1.2.")
    add_body(document, payload.get("period"))
    add_subheading(document, "Herramientas", "1.3.")
    tools = payload.get("tools") or []
    if tools:
        add_bullets(document, [f"{tool.get('name')}: {tool.get('description')}" for tool in tools])
    else:
        add_bullets(document, ["No se confirmaron herramientas específicas desde la evidencia entregada."])
    add_subheading(document, "Datos Base", "1.4.")
    add_body(document, payload.get("data_base"))
    add_subheading(document, "Cobertura del servicio", "1.5.")
    add_body(document, payload.get("coverage_summary"))
    add_coverage_table(document, payload.get("coverage_rows") or [])

    add_heading(document, "Resumen ejecutivo del dominio", "2.")
    add_body(document, payload.get("executive_summary"))
    comparison = payload.get("vulnerability_comparison") or {}
    add_subheading(document, "Distribución actual de hallazgos por severidad", "2.1.")
    add_body(document, comparison.get("summary"))
    severity_rows = comparison.get("severity_rows") or []
    if severity_rows:
        if comparison_table_template is not None:
            add_template_comparison_table(document, comparison_table_template, severity_rows)
        else:
            add_severity_comparison_table(document, severity_rows)
    add_subheading(document, "Estado actual de la seguridad", "2.2.")
    add_body(document, payload.get("histogram_summary"))
    add_subheading(document, "Focos prioritarios", "2.3.")
    add_bullets(document, payload.get("priority_focuses"))
    add_subheading(document, "Consideraciones operativas", "2.4.")
    add_bullets(document, payload.get("operational_considerations"))
    if include_admin_sections:
        add_subheading(document, "Resultados obtenidos y próximas acciones", "2.5.")
        add_body(document, payload.get("results_and_next_actions"))
        add_subheading(document, "Resultados obtenidos", "2.6.")
        add_body(document, payload.get("results_obtained"))
        add_subheading(document, "Próximas acciones", "2.7.")
        add_bullets(document, payload.get("next_actions"))
        add_subheading(document, "Requerimiento", "2.7.1.")
        add_bullets(document, payload.get("requirements"))

    add_heading(document, "Seguridad por Dominio", "3.")
    domains = payload.get("security_domains") or []
    for index, domain in enumerate(domains, start=1):
        add_subheading(document, domain.get("name") or f"Dominio {index}", f"3.{index}.")
        layer = _domain_layer(domain)
        kpis = _domain_soc_kpis(domain) if layer == "SOC" else _domain_noc_kpis(domain)
        add_domain_snapshot_table(document, kpis, fill=SOC_FILL if layer == "SOC" else NOC_FILL)
        add_body(document, domain.get("summary"))
        add_callout(document, _domain_callout_title(domain), _domain_callout_text(domain), fill=SOC_FILL if layer == "SOC" else NOC_FILL)
        if not (domain.get("findings") or []):
            add_body(document, "No se registraron hallazgos indexados para esta integración en el período evaluado.")
            continue
        template_table = (domain_table_templates or [None])[0]
        if template_table is not None:
            add_template_findings_table(document, template_table, domain.get("findings") or [], layer=layer)
        else:
            add_findings_table(document, domain.get("findings") or [], layer=layer)

    if include_admin_sections:
        add_heading(document, "Reporte de acciones trabajadas durante la semana", "4.")
        add_bullets(document, payload.get("weekly_actions"))

        add_heading(document, "Resultados obtenidos", "5.")
        add_subheading(document, "Seguridad Reforzada", "5.1.")
        add_body(document, payload.get("reinforced_security"))
        add_subheading(document, "Hallazgos pendientes", "5.2.")
        add_bullets(document, payload.get("pending_findings"))

        add_heading(document, "Noticias de seguridad", "6.")
        for index, news in enumerate(payload.get("security_news") or [], start=1):
            add_subheading(document, news.get("title") or f"Noticia {index}", f"6.{index}.")
            add_key_value_table(
                document,
                [
                    ("Fecha", news.get("date")),
                    ("Fuente", news.get("source")),
                    ("Enlaces", ", ".join(news.get("links") or [])),
                ],
            )
            add_body(document, news.get("summary"))
            add_callout(document, "Recomendación", news.get("recommendation"))

    if payload.get("limitations"):
        add_callout(document, "Limitaciones del análisis", bullets=payload.get("limitations"), fill=WARNING)


def _replace_paragraph_text(paragraph: Paragraph, value: Any) -> None:
    text = _clean_text(value)
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def _section_keys_for_variant(report_variant: str) -> list[str]:
    keys = [
        "service",
        "period",
        "tools",
        "data_base",
        "executive",
        "comparison",
        "histogram",
        "domains",
    ]
    if report_variant == "client_admin":
        keys.extend(
            [
                "results_next",
                "results_obtained",
                "next_actions",
                "requirements",
                "weekly_actions",
                "reinforced_security",
                "pending_findings",
                "security_news",
            ]
        )
    return keys


def _find_section_anchor(document: Document, key: str, *, required: bool = True) -> Paragraph | None:
    anchor = _find_paragraph_containing(document, SECTION_ANCHORS[key])
    if anchor is None and required:
        raise RuntimeError(f"Missing template anchor for section '{key}': {SECTION_ANCHORS[key][0]}")
    return anchor


def _collect_section_anchors(document: Document, report_variant: str) -> dict[str, Paragraph]:
    keys = _section_keys_for_variant(report_variant)
    if report_variant == "client_admin":
        keys.append("results")
    anchors: dict[str, Paragraph] = {}
    for key in keys:
        required = key not in {"results"}
        anchor = _find_section_anchor(document, key, required=required)
        if anchor is not None:
            anchors[key] = anchor
    return anchors


def _next_section_anchor(anchors: dict[str, Paragraph], current_key: str) -> Paragraph | None:
    current_index = SECTION_ORDER.index(current_key)
    for next_key in SECTION_ORDER[current_index + 1 :]:
        anchor = anchors.get(next_key)
        if anchor is not None:
            return anchor
    return None


def _remove_body_between(anchor: Paragraph, next_anchor: Paragraph | None) -> None:
    current = anchor._p.getnext()
    while current is not None and current is not (next_anchor._p if next_anchor is not None else None):
        following = current.getnext()
        if current.tag != qn("w:sectPr"):
            current.getparent().remove(current)
        current = following


def _insert_paragraph_after_node(document: Document, node: Any, *, style_name: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    node.addnext(new_p)
    paragraph = Paragraph(new_p, document._body)
    if style_name and style_name in document.styles:
        paragraph.style = style_name
    return paragraph


def _insert_text_after_node(
    document: Document,
    node: Any,
    text: Any,
    *,
    style_candidates: tuple[str, ...] = ("Body Text", "Normal"),
    justify: bool = True,
    empty_text: str = "Sin información confirmada para esta sección.",
) -> Any:
    chunks = [_clean_text(line) for line in str(text or "").splitlines() if _clean_text(line)]
    if not chunks:
        chunks = [empty_text]
    current = node
    for chunk in chunks:
        paragraph = _insert_paragraph_after_node(document, current, style_name=_style_name(document, *style_candidates))
        if justify:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _spacing(paragraph, after=5)
        paragraph.add_run(chunk)
        current = paragraph._p
    return current


def _insert_list_after_node(
    document: Document,
    node: Any,
    values: list[str],
    *,
    style_candidates: tuple[str, ...] = ("List Paragraph", "Body Text", "Normal"),
) -> Any:
    current = node
    entries = [_clean_text(value) for value in values if _clean_text(value)]
    if not entries:
        return _insert_text_after_node(document, current, "Sin información confirmada para esta sección.")
    for value in entries:
        paragraph = _insert_paragraph_after_node(document, current, style_name=_style_name(document, *style_candidates))
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _spacing(paragraph, after=4)
        paragraph.add_run(f"- {value}")
        current = paragraph._p
    return current


def _insert_table_after_node(document: Document, node: Any, rows: int, cols: int) -> Table:
    table = document.add_table(rows=rows, cols=cols)
    node.addnext(table._tbl)
    return table


def _insert_severity_comparison_table_after_node(document: Document, node: Any, rows: list[dict[str, Any]]) -> Any:
    if not rows:
        return node
    table = _insert_table_after_node(document, node, 1, 3)
    table.autofit = False
    widths = [1.35, 1.25, 1.25]
    try:
        table.style = "Table Normal"
    except KeyError:
        pass
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_row_widths(table.rows[0], widths)
    for cell, header in zip(table.rows[0].cells, ("Severidad", "Referencia anterior", "Estado actual")):
        _shade_cell(cell, "006D9F")
        _set_cell_border(cell, "BFBFBF", "4")
        _set_cell_text(cell, header, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=8.8)
    for row in rows:
        cells = table.add_row().cells
        _set_row_widths(table.rows[-1], widths)
        for cell in cells:
            _set_cell_border(cell, "BFBFBF", "4")
        severity = row.get("severity")
        _shade_cell(cells[0], _severity_color_fill(severity))
        _set_cell_text(cells[0], severity, bold=True, color=_severity_text_color(severity), size=8.5)
        _set_cell_text(cells[1], row.get("previous"), size=8.5)
        _set_cell_text(cells[2], row.get("current"), size=8.5)
    return table._tbl


def _insert_domain_table_after_node(document: Document, node: Any, template_table: Any, findings: list[dict[str, Any]]) -> Any:
    table = Table(deepcopy(template_table), document)
    node.addnext(table._tbl)
    _populate_template_findings_table(table, findings)
    return table._tbl


def _insert_domain_heading_after_node(document: Document, node: Any, text: str) -> Any:
    paragraph = _insert_paragraph_after_node(document, node, style_name=_style_name(document, "Heading 3", "Normal"))
    _spacing(paragraph, before=7, after=4)
    paragraph.add_run(text)
    return paragraph._p


def _render_section_text(document: Document, anchors: dict[str, Paragraph], key: str, value: Any) -> None:
    anchor = anchors[key]
    next_anchor = _next_section_anchor(anchors, key)
    _remove_body_between(anchor, next_anchor)
    _insert_text_after_node(document, anchor._p, value)


def _render_section_list(document: Document, anchors: dict[str, Paragraph], key: str, values: list[str]) -> None:
    anchor = anchors[key]
    next_anchor = _next_section_anchor(anchors, key)
    _remove_body_between(anchor, next_anchor)
    _insert_list_after_node(document, anchor._p, values)


def _render_tools_section(document: Document, anchors: dict[str, Paragraph], payload: dict[str, Any]) -> None:
    tools = [
        f"{tool.get('name') or 'Herramienta'}: {tool.get('description') or ''}".strip()
        for tool in (payload.get("tools") or [])
        if isinstance(tool, dict)
    ]
    _render_section_list(document, anchors, "tools", tools)


def _render_comparison_section(document: Document, anchors: dict[str, Paragraph], payload: dict[str, Any]) -> None:
    anchor = anchors["comparison"]
    next_anchor = _next_section_anchor(anchors, "comparison")
    _remove_body_between(anchor, next_anchor)
    comparison = payload.get("vulnerability_comparison") or {}
    current = _insert_text_after_node(document, anchor._p, comparison.get("summary"))
    current = _insert_severity_comparison_table_after_node(document, current, comparison.get("severity_rows") or [])
    _insert_paragraph_after_node(document, current)


def _render_domains_section(
    document: Document,
    anchors: dict[str, Paragraph],
    payload: dict[str, Any],
    domain_table_templates: list[Any] | None,
) -> None:
    anchor = anchors["domains"]
    next_anchor = _next_section_anchor(anchors, "domains")
    _remove_body_between(anchor, next_anchor)
    current: Any = anchor._p
    domains = payload.get("security_domains") or []
    if not domains:
        _insert_text_after_node(document, current, "Sin dominios de seguridad confirmados para el período evaluado.")
        return
    for index, domain in enumerate(domains, start=1):
        title = domain.get("name") or f"Dominio {index}"
        current = _insert_domain_heading_after_node(document, current, f"3.{index}. {title}")
        current = _insert_text_after_node(document, current, domain.get("summary"))
        current = _insert_text_after_node(document, current, f"{_domain_callout_title(domain)}: {_domain_callout_text(domain)}")
        if not (domain.get("findings") or []):
            current = _insert_text_after_node(document, current, "No se registraron hallazgos indexados para esta integración en el período evaluado.")
            current = _insert_paragraph_after_node(document, current)._p
            continue
        template_table = (domain_table_templates or [])[index - 1] if index <= len(domain_table_templates or []) else None
        findings = domain.get("findings") or []
        if template_table is not None:
            current = _insert_domain_table_after_node(document, current, template_table, findings)
        else:
            table = document.add_table(rows=1, cols=4)
            current.addnext(table._tbl)
            _populate_template_findings_table(table, findings)
            current = table._tbl
        current = _insert_paragraph_after_node(document, current)._p


def _render_news_section(document: Document, anchors: dict[str, Paragraph], payload: dict[str, Any]) -> None:
    anchor = anchors["security_news"]
    next_anchor = _next_section_anchor(anchors, "security_news")
    _remove_body_between(anchor, next_anchor)
    current: Any = anchor._p
    news_items = payload.get("security_news") or []
    if not news_items:
        _insert_text_after_node(document, current, "Sin noticias de seguridad confirmadas para este reporte.")
        return
    for index, news in enumerate(news_items, start=1):
        current = _insert_domain_heading_after_node(document, current, f"6.{index}. {news.get('title') or f'Noticia {index}'}")
        current = _insert_text_after_node(document, current, f"Fecha: {_clean_text(news.get('date')) or 'No especificada'}", justify=False)
        current = _insert_text_after_node(document, current, f"Fuente: {_clean_text(news.get('source')) or 'No especificada'}", justify=False)
        links = ", ".join(news.get("links") or [])
        if links:
            current = _insert_text_after_node(document, current, f"Enlaces: {links}", justify=False)
        current = _insert_text_after_node(document, current, news.get("summary"))
        recommendation = _clean_text(news.get("recommendation"))
        if recommendation:
            current = _insert_text_after_node(document, current, f"Recomendación: {recommendation}")
        current = _insert_paragraph_after_node(document, current)._p


def _section_has_visible_content(anchor: Paragraph, next_anchor: Paragraph | None) -> bool:
    current = anchor._p.getnext()
    end = next_anchor._p if next_anchor is not None else None
    while current is not None and current is not end:
        text = "".join(node.text or "" for node in current.iter(qn("w:t"))).strip()
        if text:
            return True
        if current.tag == qn("w:tbl"):
            return True
        current = current.getnext()
    return False


def assert_required_sections_filled(document: Document, anchors: dict[str, Paragraph], report_variant: str) -> None:
    empty_sections: list[str] = []
    for key in _section_keys_for_variant(report_variant):
        anchor = anchors.get(key)
        if anchor is None:
            empty_sections.append(key)
            continue
        if not _section_has_visible_content(anchor, _next_section_anchor(anchors, key)):
            empty_sections.append(key)
    if empty_sections:
        raise RuntimeError(f"Rendered sections are empty or missing: {', '.join(empty_sections)}")


def _paragraph_after_heading(document: Document, heading: str, occurrence: int = 0) -> Paragraph | None:
    matches = [p for p in document.paragraphs if _plain_match_text(heading) in _plain_match_text(p.text)]
    if occurrence >= len(matches):
        return None
    return _next_content_paragraph(document, matches[occurrence])


def _clear_demo_text(document: Document) -> None:
    """Vacía únicamente contenido de ejemplo; conserva la estructura Word."""
    markers = (
        "lorem",
        "oreм ipsum",
        "change for date",
        "change for prepared for",
        "change date",
        "change.",
        "change",
    )
    for paragraph in document.paragraphs:
        normalized = _plain_match_text(paragraph.text)
        if any(marker in normalized for marker in markers):
            _replace_paragraph_text(paragraph, "")
    for box in document._element.xpath(".//w:txbxContent"):
        original = _plain_match_text(_textbox_text(box).strip())
        if any(marker in original for marker in markers):
            _clear_textbox_content(box)


def _template_residue_matches(value: str) -> bool:
    normalized = _plain_match_text(value)
    return any(marker in normalized for marker in (
        "lorem",
        "change for tenant",
        "change for",
        "change for date",
        "change for prepared for",
        "change date",
        "change.",
        "oreм ipsum",
    ))


def assert_no_template_residue(document: Document) -> None:
    leftovers: list[str] = []
    for paragraph in document.paragraphs:
        text = (paragraph.text or "").strip()
        if text and _template_residue_matches(text):
            leftovers.append(text)
    for box in document._element.xpath(".//w:txbxContent"):
        text = _textbox_text(box).strip()
        if text and _template_residue_matches(text):
            leftovers.append(text)
    for section in document.sections:
        for story in (section.header, section.footer):
            for paragraph in story.paragraphs:
                text = (paragraph.text or "").strip()
                if text and _template_residue_matches(text):
                    leftovers.append(text)
            for box in story._element.xpath(".//w:txbxContent"):
                text = _textbox_text(box).strip()
                if text and _template_residue_matches(text):
                    leftovers.append(text)
    if leftovers:
        sample = " | ".join(leftovers[:5])
        raise RuntimeError(f"Template residue detected after fill: {sample}")


def _fill_existing_table_slots(document: Document, payload: dict[str, Any]) -> None:
    domains = payload.get("security_domains") or []
    tables = [table for table in document.tables if len(table.columns) == 4 and table.rows]
    for index, table in enumerate(tables):
        domain = domains[index] if index < len(domains) else {}
        _populate_template_findings_table(table, domain.get("findings") or [])


def _fill_existing_action_slots(document: Document, heading: str, values: list[str]) -> None:
    anchor = _find_paragraph_containing(document, [heading])
    if anchor is None:
        return
    start = False
    slots: list[Paragraph] = []
    for paragraph in document.paragraphs:
        if paragraph._p is anchor._p:
            start = True
            continue
        if not start:
            continue
        if paragraph.text.strip() and paragraph.style.name.startswith("Heading"):
            break
        if not paragraph.text.strip() or "lorem" in _plain_match_text(paragraph.text):
            slots.append(paragraph)
    for index, slot in enumerate(slots):
        _replace_paragraph_text(slot, values[index] if index < len(values) else "")


def fill_minority_template_in_place(document: Document, payload: dict[str, Any]) -> None:
    """Rellena la plantilla corporativa por secciones ancladas y validables."""
    _clear_demo_text(document)
    report_variant = normalize_report_variant(
        payload.get("report_variant")
        or payload.get("template_variant")
        or (payload.get("mock_meta") or {}).get("report_variant")
    )
    anchors = _collect_section_anchors(document, report_variant)
    domain_table_templates = capture_domain_table_templates(document)

    service_text = (
        f"{payload.get('service_name') or 'Servicio de monitoreo proactivo XOC'} implementado por "
        f"{payload.get('prepared_by') or 'TXDXSECURE'} para el cliente {payload.get('client_name') or 'Cliente'}."
    )
    _render_section_text(document, anchors, "service", service_text)
    _render_section_text(document, anchors, "period", payload.get("period"))
    _render_tools_section(document, anchors, payload)
    _render_section_text(document, anchors, "data_base", payload.get("data_base"))
    _render_section_text(document, anchors, "executive", payload.get("executive_summary"))
    _render_comparison_section(document, anchors, payload)
    _render_section_text(document, anchors, "histogram", payload.get("histogram_summary"))
    _render_domains_section(document, anchors, payload, domain_table_templates)

    if report_variant == "client_admin":
        _render_section_text(document, anchors, "results_next", payload.get("results_and_next_actions"))
        _render_section_text(document, anchors, "results_obtained", payload.get("results_obtained"))
        _render_section_list(document, anchors, "next_actions", payload.get("next_actions") or [])
        _render_section_list(document, anchors, "requirements", payload.get("requirements") or [])
        _render_section_list(document, anchors, "weekly_actions", payload.get("weekly_actions") or [])
        _render_section_text(document, anchors, "reinforced_security", payload.get("reinforced_security"))
        _render_section_list(document, anchors, "pending_findings", payload.get("pending_findings") or [])
        _render_news_section(document, anchors, payload)

    assert_required_sections_filled(document, anchors, report_variant)


def _insert_paragraph_after(paragraph: Paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def _caption_after(paragraph: Paragraph, text: str) -> Paragraph:
    caption = _insert_paragraph_after(paragraph)
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_after = Pt(8)
    run = caption.add_run(text)
    run.italic = True
    run.font.size = Pt(8.5)
    run.font.color.rgb = MUTED_TEXT
    return caption


def _plain_match_text(value: str) -> str:
    normalized = unicodedata.normalize("NFKD", value or "")
    return "".join(char for char in normalized if not unicodedata.combining(char)).lower()


def _find_paragraph_containing(document: Document, patterns: list[str]) -> Paragraph | None:
    normalized_patterns = [_plain_match_text(pattern) for pattern in patterns]
    for paragraph in document.paragraphs:
        text = _plain_match_text(paragraph.text)
        if any(pattern in text for pattern in normalized_patterns):
            return paragraph
    return None


def _find_last_paragraph_containing(document: Document, patterns: list[str]) -> Paragraph | None:
    normalized_patterns = [_plain_match_text(pattern) for pattern in patterns]
    for paragraph in reversed(document.paragraphs):
        text = _plain_match_text(paragraph.text)
        if any(pattern in text for pattern in normalized_patterns):
            return paragraph
    return None


def _next_content_paragraph(document: Document, anchor: Paragraph) -> Paragraph | None:
    found = False
    for paragraph in document.paragraphs:
        if paragraph._p is anchor._p:
            found = True
            continue
        if not found:
            continue
        if paragraph.text.strip():
            return paragraph
    return None


def _insert_pie_pair_after(anchor: Paragraph, images: list[Path], descriptions: list[str]) -> Paragraph:
    tag_1 = _insert_paragraph_after(anchor)
    tag_1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    tag_1.paragraph_format.space_before = Pt(6)
    tag_1.paragraph_format.space_after = Pt(2)
    run_1 = tag_1.add_run("FIGURA 1")
    _format_run(run_1, bold=True, color=BRAND_BLUE, size=8)

    picture_1 = _insert_paragraph_after(tag_1)
    picture_1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture_1.paragraph_format.space_before = Pt(0)
    picture_1.paragraph_format.space_after = Pt(0)
    picture_1.add_run().add_picture(str(images[0]), width=Inches(4.15))
    desc_1 = descriptions[0] if len(descriptions) > 0 else "Referencia anterior disponible."
    caption_1 = _caption_after(picture_1, f"Figura 1. {desc_1}")

    tag_2 = _insert_paragraph_after(caption_1)
    tag_2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    tag_2.paragraph_format.space_before = Pt(10)
    tag_2.paragraph_format.space_after = Pt(2)
    run_2 = tag_2.add_run("FIGURA 2")
    _format_run(run_2, bold=True, color=BRAND_BLUE, size=8)

    picture_2 = _insert_paragraph_after(tag_2)
    picture_2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture_2.paragraph_format.space_before = Pt(0)
    picture_2.paragraph_format.space_after = Pt(0)
    picture_2.add_run().add_picture(str(images[1]), width=Inches(4.15))
    desc_2 = descriptions[1] if len(descriptions) > 1 else "Estado actual."
    return _caption_after(picture_2, f"Figura 2. {desc_2}")


def place_evidence_images(document: Document, images: list[Path], descriptions: list[str]) -> None:
    if not images:
        return
    fallback = _find_last_paragraph_containing(document, ["Resumen ejecutivo", "Seguridad por Dominio"]) or document.paragraphs[-1]
    inserted_after = fallback
    last_citation_anchor_el = None
    skipped_indexes: set[int] = set()
    if len(images) >= 2:
        pie_anchor = _find_last_paragraph_containing(document, ["Distribución actual de hallazgos por severidad"])
        if pie_anchor is not None:
            inserted_after = _insert_pie_pair_after(_next_content_paragraph(document, pie_anchor) or pie_anchor, images[:2], descriptions[:2])
            skipped_indexes.update({1, 2})
    for index, path in enumerate(images, start=1):
        if index in skipped_indexes:
            continue
        label = f"Figura {index}"
        target = None
        for paragraph in document.paragraphs:
            if label.lower() in paragraph.text.lower():
                target = paragraph
                break
        if target is None:
            if index in {1, 2}:
                target = _find_last_paragraph_containing(document, ["Distribución actual de hallazgos por severidad"])
            elif index == 3:
                target = _find_last_paragraph_containing(document, ["Estado actual de la seguridad"])
        citation_anchor_el = target._p if target is not None else None
        if target is None or citation_anchor_el is last_citation_anchor_el:
            target = inserted_after
        if citation_anchor_el is not None:
            last_citation_anchor_el = citation_anchor_el
        tag = _insert_paragraph_after(target)
        tag.alignment = WD_ALIGN_PARAGRAPH.LEFT
        tag.paragraph_format.space_before = Pt(8)
        tag.paragraph_format.space_after = Pt(3)
        run = tag.add_run(label.upper())
        _format_run(run, bold=True, color=BRAND_BLUE, size=8)
        picture = _insert_paragraph_after(tag)
        picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
        picture.paragraph_format.space_before = Pt(0)
        picture.paragraph_format.space_after = Pt(0)
        image_width = 4.15 if index in {1, 2} else 5.7
        picture.add_run().add_picture(str(path), width=Inches(image_width))
        description = descriptions[index - 1] if index - 1 < len(descriptions) else ""
        caption = _caption_after(picture, f"{label}. {description or 'Evidencia visual proporcionada.'}")
        inserted_after = caption


def validate_docx(path: Path) -> None:
    if not path.exists() or path.stat().st_size < 10_000:
        raise RuntimeError("El DOCX no fue generado correctamente")
    with zipfile.ZipFile(path) as archive:
        if "word/document.xml" not in archive.namelist():
            raise RuntimeError("El DOCX generado no contiene word/document.xml")
        leftovers: list[str] = []
        for name in archive.namelist():
            if not name.startswith("word/") or not name.endswith(".xml"):
                continue
            try:
                content = archive.read(name).decode("utf-8", errors="ignore")
            except KeyError:
                continue
            if _template_residue_matches(content):
                leftovers.append(name)
        if leftovers:
            raise RuntimeError(f"El DOCX generado conserva residuos de plantilla en: {', '.join(leftovers[:5])}")


def _chart_images_from_payload(payload: dict[str, Any]) -> tuple[list[Path], list[str]]:
    images: list[Path] = []
    descriptions: list[str] = []
    for item in payload.get("chart_images") or []:
        if not isinstance(item, dict):
            continue
        path_value = _clean_text(item.get("path"))
        if not path_value:
            continue
        path = Path(path_value)
        if path.is_file():
            images.append(path)
            descriptions.append(_clean_text(item.get("description")))
    return images, descriptions


def generate_minority_report_docx(template_path: str | None, payload: dict[str, Any], output_path: str) -> str:
    """Genera el DOCX del Minority Report desde una plantilla local/S3.

    `template_path` se respeta si viene del storage del backend. Si viene vacío o
    no existe, se escoge la plantilla empaquetada según `report_variant`.
    """
    report_variant = normalize_report_variant(
        payload.get("report_variant")
        or payload.get("template_variant")
        or (payload.get("mock_meta") or {}).get("report_variant")
    )
    selected_template = Path(template_path) if template_path and Path(template_path).is_file() else base_template_path(report_variant)
    result_path = Path(output_path)
    result_path.parent.mkdir(parents=True, exist_ok=True)

    document = Document(selected_template)
    apply_example_body_style(document)
    update_cover_and_footer(document, payload)
    domain_table_templates = capture_domain_table_templates(document)
    comparison_table_template = capture_comparison_table_template(document)
    clear_template_body_after_cover(document)
    build_report_body(document, payload, domain_table_templates, comparison_table_template)
    images, descriptions = _chart_images_from_payload(payload)
    place_evidence_images(document, images, descriptions)
    assert_no_template_residue(document)

    document.save(result_path)
    validate_docx(result_path)
    return str(result_path)


def _read_optional_text(path: Path | None) -> str:
    if not path:
        return ""
    return path.read_text(encoding="utf-8")


def _read_optional_json(path: Path | None) -> dict[str, Any]:
    if not path:
        return {}
    data = json.loads(path.read_text(encoding="utf-8"))
    if not isinstance(data, dict):
        raise ValueError("El JSON estructurado debe ser un objeto.")
    return data


def _text_from_args(args: argparse.Namespace) -> str:
    pieces = []
    if args.text:
        pieces.append(args.text)
    if args.text_file:
        pieces.append(args.text_file.read_text(encoding="utf-8"))
    return "\n\n".join(piece.strip() for piece in pieces if piece.strip())


def main() -> None:
    parser = argparse.ArgumentParser(description="Genera Minority Report XOC usando Azure Foundry multimodal.")
    parser.add_argument("--client-name", required=True, help="Cliente preparado para portada y footer.")
    parser.add_argument("--period", required=True, help="Periodo del reporte. Ejemplo: Del 20 al 26 de junio de 2026.")
    parser.add_argument("--text", default="", help="Texto libre del analista.")
    parser.add_argument("--text-file", type=Path, help="Archivo TXT/MD con contexto del analista.")
    parser.add_argument("--data", type=Path, help="JSON estructurado opcional.")
    parser.add_argument("--reference-md", type=Path, help="Markdown de referencia del Minority Report ejemplo.")
    parser.add_argument("--image", action="append", default=[], type=Path, help="Ruta de imagen local. Se puede repetir.")
    parser.add_argument("--image-description", action="append", default=[], help="Descripción de imagen. Repetir en el mismo orden que --image.")
    parser.add_argument(
        "--report-variant",
        choices=["client", "client_admin"],
        default="client",
        help="Plantilla Minority a usar: report for client o report for admin client.",
    )
    parser.add_argument("--no-azure", action="store_true", help="No llama Azure; genera borrador local.")
    parser.add_argument("--allow-local-fallback", action="store_true", default=True, help="Si Azure falla, genera borrador local.")
    args = parser.parse_args()
    load_dotenv(ROOT / ".env")

    try:
        model_path = base_template_path(args.report_variant)
    except FileNotFoundError as exc:
        raise SystemExit(str(exc)) from exc

    text = _text_from_args(args)
    structured = _read_optional_json(args.data)
    reference_markdown = _read_optional_text(args.reference_md)
    payload = generate_minority_payload(
        client_name=args.client_name,
        period=args.period,
        analyst_text=text,
        image_paths=args.image,
        image_descriptions=args.image_description,
        structured_data=structured,
        reference_markdown=reference_markdown,
        use_azure=not args.no_azure,
        allow_local_fallback=args.allow_local_fallback,
        output_path=DEFAULT_PAYLOAD_PATH,
    )
    payload["report_variant"] = args.report_variant
    payload["template_variant"] = "report for admin client" if args.report_variant == "client_admin" else "report for client"

    document = Document(model_path)
    apply_example_body_style(document)
    update_cover_and_footer(document, payload)
    domain_table_templates = capture_domain_table_templates(document)
    comparison_table_template = capture_comparison_table_template(document)
    clear_template_body_after_cover(document)
    build_report_body(document, payload, domain_table_templates, comparison_table_template)
    place_evidence_images(document, args.image, args.image_description)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    result = _writable_output_path(_output_path(payload))
    document.save(result)
    validate_docx(result)
    print(f"Payload JSON: {DEFAULT_PAYLOAD_PATH}")
    print(f"Minority Report generado: {result}")


if __name__ == "__main__":
    main()

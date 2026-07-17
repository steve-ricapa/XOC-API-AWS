from __future__ import annotations

import hashlib
import json
import os
import zipfile
from collections import Counter
from copy import deepcopy
from datetime import datetime
from pathlib import Path
from typing import Any

import matplotlib
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from lxml import etree

from azure_foundry_agent_client import generate_ai_sections
from database_report_loader import build_report_data_from_database

try:
    from dotenv import load_dotenv
except ImportError:  # pragma: no cover - optional local dependency during bootstrap
    def load_dotenv() -> bool:
        return False

matplotlib.use("Agg")
import matplotlib.pyplot as plt


load_dotenv()

BASE_DIR = Path(__file__).resolve().parent
DATA_DIR = BASE_DIR / "data"
TEMPLATES_DIR = BASE_DIR / "templates"
OUTPUT_DIR = BASE_DIR / "output"
REPORTS_ROOT_DIR = OUTPUT_DIR / "reports"

MOCK_DATA_PATH = DATA_DIR / "mock_report_data.json"
TEMPLATE_PATH = TEMPLATES_DIR / "security-report-v1.docx"
DATABASE_REPORT_DATA_PATH = OUTPUT_DIR / "report_data_from_database.json"
AZURE_SECTIONS_OUTPUT_PATH = OUTPUT_DIR / "azure_foundry_generated_sections.json"

SEVERITY_CHART_PATH = OUTPUT_DIR / "severity_comparison_chart.png"
DOMAIN_SCORE_CHART_PATH = OUTPUT_DIR / "domain_score_chart.png"
CASES_BY_SEVERITY_CHART_PATH = OUTPUT_DIR / "cases_by_severity_chart.png"
SLA_STATUS_CHART_PATH = OUTPUT_DIR / "sla_status_chart.png"

DEFAULT_MOCK_DATA: dict[str, Any] = {
    "tenant": {
        "id": "tenant-jockey-salud",
        "name": "Jockey Salud",
        "tenant_id": "7aa8a350-3b80-49c6-8aa7-xoc-jockey-001",
    },
    "report": {
        "id": "report-demo-001",
        "title": "Reporte Operativo Semanal XOC | Jockey Salud",
        "service": "Servicio de Monitoreo Proactivo XOC",
        "generated_at": "2026-06-26 18:00:00",
        "period": "Del 20 de junio al 26 de junio del 2026",
        "prepared_by": "TXDXSECURE",
        "scope": "Web externo, IPs publicas, firewall, servidores, switching, WIFI, desktops y activos OT/IoT priorizados por Jockey Salud.",
    },
    "tools": [
        "MonEvents",
        "MonVulE",
        "MonVulC",
        "MonApps",
        "MonNet",
        "MonInfra",
        "OpenVAS",
        "InsightVM",
        "Wazuh SIEM",
        "Zabbix",
        "Nessus",
        "Uptime Kuma",
    ],
    "severity_summary": {
        "previous": {"critical": 3, "high": 9, "medium": 27, "low": 14, "informational": 142},
        "current": {"critical": 2, "high": 7, "medium": 33, "low": 8, "informational": 147},
    },
    "security_posture": {
        "global_score": 72,
        "risk_level": "Alto",
        "trend": "Mejora parcial con riesgo alto residual",
        "critical_open": 2,
        "high_open": 7,
        "sla_at_risk": 2,
        "mttr_hours": 26,
        "main_risks": [
            "Virtualizacion fuera de soporte en infraestructura critica.",
            "Exposicion de componentes web y criptografia debil en activos publicados.",
            "Credenciales por defecto y debilidades SSH en equipos de red.",
            "Servicios heredados sobre activos OT/IoT y desktop con deuda de hardening.",
        ],
    },
    "cases": [
        {"id": "CASE-001", "title": "ESXi fuera de soporte", "severity": "Critico", "status": "Abierto", "domain": "Infraestructura de Computo - Servers", "asset": "esxi-prd-01", "created_at": "2026-06-20", "updated_at": "2026-06-26", "sla_target_hours": 72, "elapsed_hours": 44, "sla_status": "Dentro SLA", "owner": "Cliente", "summary": "Hipervisor principal fuera de soporte.", "next_action": "Planificar upgrade o migracion"},
        {"id": "CASE-002", "title": "Token expuesto en recurso web", "severity": "Alto", "status": "En analisis", "domain": "Web Externo", "asset": "portal.jockeysalud.pe", "created_at": "2026-06-20", "updated_at": "2026-06-26", "sla_target_hours": 96, "elapsed_hours": 40, "sla_status": "Dentro SLA", "owner": "XOC", "summary": "Exposicion de token de autorizacion en recurso JavaScript.", "next_action": "Rotar secreto y remover exposicion"},
        {"id": "CASE-003", "title": "Credenciales por defecto OT", "severity": "Critico", "status": "En remediacion", "domain": "Infraestructura OT/IoT", "asset": "ot-firebird-01", "created_at": "2026-06-21", "updated_at": "2026-06-26", "sla_target_hours": 48, "elapsed_hours": 45, "sla_status": "En riesgo", "owner": "Cliente", "summary": "Credenciales default sobre servicio industrial.", "next_action": "Rotar credenciales y aislar servicio"},
        {"id": "CASE-004", "title": "Cipher suites debiles en switches", "severity": "Alto", "status": "Abierto", "domain": "Infraestructura de Red - Switches", "asset": "switch-core-02", "created_at": "2026-06-21", "updated_at": "2026-06-26", "sla_target_hours": 96, "elapsed_hours": 58, "sla_status": "En riesgo", "owner": "Cliente", "summary": "Persisten algoritmos debiles sobre interfaces de administracion.", "next_action": "Actualizar baseline criptografica"},
        {"id": "CASE-005", "title": "TLS obsoleto en AP", "severity": "Medio", "status": "Abierto", "domain": "Infraestructura de Red - WIFI", "asset": "ap-core-01", "created_at": "2026-06-22", "updated_at": "2026-06-26", "sla_target_hours": 120, "elapsed_hours": 33, "sla_status": "Dentro SLA", "owner": "XOC", "summary": "TLS 1.0/1.1 habilitado en portal AP.", "next_action": "Actualizar configuracion TLS"},
        {"id": "CASE-006", "title": "Hardening parcial en desktops", "severity": "Medio", "status": "Resuelto", "domain": "Infraestructura de Computo - Desktops", "asset": "42 endpoints", "created_at": "2026-06-20", "updated_at": "2026-06-25", "sla_target_hours": 168, "elapsed_hours": 52, "sla_status": "Dentro SLA", "owner": "Cliente", "summary": "Ajustes sobre debilidades TLS y SSH en estaciones.", "next_action": "Monitorear regresion"},
    ],
    "sla_summary": {},
    "mitre_mapping": [
        {"tactic": "Initial Access", "technique": "Public-Facing Application", "evidence": "Token expuesto en recurso web", "domain": "Web Externo", "severity": "Alto", "status": "Abierto"},
        {"tactic": "Credential Access", "technique": "Default Credentials", "evidence": "Credenciales Firebird por defecto", "domain": "Infraestructura OT/IoT", "severity": "Critico", "status": "En remediacion"},
        {"tactic": "Defense Evasion", "technique": "Impair Defenses", "evidence": "Configuraciones criptograficas debiles", "domain": "Infraestructura de Red - Switches", "severity": "Alto", "status": "Abierto"},
        {"tactic": "Discovery", "technique": "Network Service Discovery", "evidence": "Timestamps e ICMP replies expuestos", "domain": "IPs Publicas", "severity": "Bajo", "status": "Monitoreado"},
        {"tactic": "Persistence", "technique": "External Remote Services", "evidence": "Servicios administrativos heredados", "domain": "Infraestructura de Computo - Servers", "severity": "Medio", "status": "Abierto"},
    ],
    "top_assets": [
        {"asset": "esxi-prd-01", "type": "Server", "domain": "Infraestructura de Computo - Servers", "risk_score": 94, "critical": 1, "high": 0, "medium": 2, "low": 0, "recommendation": "Actualizar plataforma y reforzar TLS"},
        {"asset": "portal.jockeysalud.pe", "type": "Web", "domain": "Web Externo", "risk_score": 90, "critical": 0, "high": 2, "medium": 2, "low": 0, "recommendation": "Eliminar exposicion de secretos y endurecer HTTPS"},
        {"asset": "switch-core-02", "type": "Switch", "domain": "Infraestructura de Red - Switches", "risk_score": 89, "critical": 0, "high": 2, "medium": 6, "low": 3, "recommendation": "Actualizar baseline SSH/TLS"},
        {"asset": "ot-firebird-01", "type": "OT/IoT", "domain": "Infraestructura OT/IoT", "risk_score": 91, "critical": 1, "high": 1, "medium": 4, "low": 0, "recommendation": "Cambiar credenciales y segmentar acceso"},
        {"asset": "ap-core-01", "type": "Access Point", "domain": "Infraestructura de Red - WIFI", "risk_score": 78, "critical": 0, "high": 1, "medium": 5, "low": 2, "recommendation": "Deshabilitar TLS heredado y SSH debil"},
    ],
    "domain_scores": [
        {"domain": "Web Externo", "score": 66, "risk": "Alto", "trend": "Estable", "comment": "Persisten componentes web con exposicion de secretos y HTTPS mejorable."},
        {"domain": "IPs Publicas", "score": 78, "risk": "Medio", "trend": "Estable", "comment": "Se mantienen hallazgos informativos de reconocimiento externo."},
        {"domain": "FW", "score": 81, "risk": "Medio", "trend": "Mejora", "comment": "Perimetro estable con hallazgos menores residuales."},
        {"domain": "Infraestructura de Computo - Servers", "score": 58, "risk": "Alto", "trend": "Empeora", "comment": "La criticidad se concentra en virtualizacion EOL."},
        {"domain": "Infraestructura de Red - Switches", "score": 61, "risk": "Alto", "trend": "Estable", "comment": "Riesgo sostenido por hardening criptografico insuficiente."},
        {"domain": "Infraestructura de Red - WIFI", "score": 69, "risk": "Medio", "trend": "Estable", "comment": "Persisten debilidades TLS/SSH sobre APs."},
        {"domain": "Infraestructura de Computo - Desktops", "score": 73, "risk": "Medio", "trend": "Mejora", "comment": "Se redujo parte del backlog tecnico en estaciones."},
        {"domain": "Infraestructura OT/IoT", "score": 55, "risk": "Alto", "trend": "Estable", "comment": "Se mantienen riesgos criticos por credenciales default y cifrado debil."},
    ],
    "domains": [
        {"name": "Web Externo", "summary": "El dominio web externo concentra exposicion de secretos, cipher suites vulnerables y configuraciones HTTPS que requieren ajuste inmediato.", "score": 66, "risk": "Alto", "findings": ["WE01", "WE03", "WE04", "WE05"]},
        {"name": "IPs Publicas", "summary": "Las IPs publicas mantienen hallazgos de bajo impacto, principalmente utiles para reconocimiento y enumeracion externa.", "score": 78, "risk": "Medio", "findings": ["PB01", "PB02"]},
        {"name": "FW", "summary": "El perimetro presenta exposiciones informativas residuales sin criticidad alta en este corte.", "score": 81, "risk": "Medio", "findings": ["FW01", "FW02"]},
        {"name": "Infraestructura de Computo - Servers", "summary": "Los servidores concentran el mayor impacto operativo por EOL en virtualizacion y certificados no confiables.", "score": 58, "risk": "Alto", "findings": ["SR-01", "SR-02", "SR-03", "SR-04", "SR-07", "SR-09"]},
        {"name": "Infraestructura de Red - Switches", "summary": "El dominio de switches mantiene riesgo alto por traversal, suites debiles, SSH inseguro y exposiciones informativas de red.", "score": 61, "risk": "Alto", "findings": ["SW-01", "SW-02", "SW-03", "SW-04", "SW-06"]},
        {"name": "Infraestructura de Red - WIFI", "summary": "Los APs mantienen exposicion media por SSH heredado, TLS obsoleto y algoritmos MAC debiles.", "score": 69, "risk": "Medio", "findings": ["AP01", "AP02", "AP03", "AP05"]},
        {"name": "Infraestructura de Computo - Desktops", "summary": "Las estaciones muestran mejora parcial, aunque persisten debilidades TLS, SSH y enumeracion RPC.", "score": 73, "risk": "Medio", "findings": ["DK-01", "DK-02", "DK-03", "DK-04"]},
        {"name": "Infraestructura OT/IoT", "summary": "Los activos OT/IoT mantienen el riesgo mas sensible por credenciales default y controles criptograficos heredados.", "score": 55, "risk": "Alto", "findings": ["OT-01", "OT-02", "OT-03", "OT-04"]},
    ],
    "findings": [
        {"id": "WE01", "domain": "Web Externo", "title": "Exposicion de token de autorizacion en recurso JavaScript publico", "affected_hosts": "1", "severity": "Alto", "description": "Se identifico un token expuesto en un recurso publico.", "recommendation": "Rotar secreto y remover referencia del codigo cliente.", "evidence": "Recurso JavaScript accesible desde Internet con token embebido.", "source_tool": "MonApps", "status": "Pendiente", "remediation_priority": "Alta"},
        {"id": "WE03", "domain": "Web Externo", "title": "SSL/TLS vulnerable cipher suites for HTTPS", "affected_hosts": "7", "severity": "Alto", "description": "Suites criptograficas no recomendadas en servicios HTTPS.", "recommendation": "Restringir a suites modernas y revisar compatibilidad.", "evidence": "Handshake HTTPS con suites vulnerables.", "source_tool": "Nessus", "status": "Pendiente", "remediation_priority": "Alta"},
        {"id": "WE04", "domain": "Web Externo", "title": "SSL/TLS weak cipher suites", "affected_hosts": "6", "severity": "Medio", "description": "Persisten cipher suites debiles.", "recommendation": "Endurecer configuracion TLS.", "evidence": "Enumeracion criptografica insegura.", "source_tool": "OpenVAS", "status": "Pendiente", "remediation_priority": "Media"},
        {"id": "WE05", "domain": "Web Externo", "title": "Certificado con claves RSA menores a 2048 bits", "affected_hosts": "7", "severity": "Medio", "description": "Cadena de certificados con fortaleza insuficiente.", "recommendation": "Renovar certificados con claves robustas.", "evidence": "Cadena TLS con claves RSA debiles.", "source_tool": "Nessus", "status": "Pendiente", "remediation_priority": "Media"},
        {"id": "PB01", "domain": "IPs Publicas", "title": "TCP timestamps information disclosure", "affected_hosts": "1", "severity": "Bajo", "description": "Activos expuestos responden timestamps TCP.", "recommendation": "Reducir informacion util para fingerprinting.", "evidence": "Timestamp TCP visible.", "source_tool": "MonNet", "status": "Pendiente", "remediation_priority": "Baja"},
        {"id": "PB02", "domain": "IPs Publicas", "title": "ICMP timestamp reply information disclosure", "affected_hosts": "1", "severity": "Bajo", "description": "Respuesta ICMP con informacion de tiempo.", "recommendation": "Restringir respuestas ICMP no necesarias.", "evidence": "ICMP timestamp reply habilitado.", "source_tool": "MonNet", "status": "Pendiente", "remediation_priority": "Baja"},
        {"id": "FW01", "domain": "FW", "title": "TCP timestamps information disclosure", "affected_hosts": "1", "severity": "Bajo", "description": "Timestamp TCP visible sobre perimetro.", "recommendation": "Limitar exposicion informativa.", "evidence": "Respuesta TCP con timestamp.", "source_tool": "MonInfra", "status": "Pendiente", "remediation_priority": "Baja"},
        {"id": "FW02", "domain": "FW", "title": "ICMP timestamp reply information disclosure", "affected_hosts": "1", "severity": "Bajo", "description": "Respuestas ICMP informativas sobre perimetro.", "recommendation": "Ajustar politicas de respuesta.", "evidence": "ICMP timestamp response.", "source_tool": "MonInfra", "status": "Pendiente", "remediation_priority": "Baja"},
        {"id": "SR-01", "domain": "Infraestructura de Computo - Servers", "title": "Operating System End of Life Detection (ESXi)", "affected_hosts": "1", "severity": "Critico", "description": "Plataforma ESXi fuera de soporte.", "recommendation": "Actualizar o migrar carga critica.", "evidence": "Version EOL detectada por escaneo.", "source_tool": "Nessus", "status": "Pendiente", "remediation_priority": "Critica"},
        {"id": "SR-02", "domain": "Infraestructura de Computo - Servers", "title": "SSL Certificate Cannot Be Trusted", "affected_hosts": "15", "severity": "Medio", "description": "Certificados no confiables en servidores.", "recommendation": "Renovar certificados validos.", "evidence": "Cadenas TLS no confiables.", "source_tool": "OpenVAS", "status": "Pendiente", "remediation_priority": "Media"},
        {"id": "SR-03", "domain": "Infraestructura de Computo - Servers", "title": "SSL Self-Signed Certificate", "affected_hosts": "15", "severity": "Medio", "description": "Certificados autofirmados presentes.", "recommendation": "Reemplazar por certificados emitidos por CA valida.", "evidence": "Servicios con certificados self-signed.", "source_tool": "OpenVAS", "status": "Pendiente", "remediation_priority": "Media"},
        {"id": "SR-04", "domain": "Infraestructura de Computo - Servers", "title": "SSL Certificate with Wrong Hostname", "affected_hosts": "4", "severity": "Medio", "description": "Hostname no coincide con certificado.", "recommendation": "Corregir emision y bindings.", "evidence": "Mismatch entre CN/SAN y host.", "source_tool": "Nessus", "status": "Pendiente", "remediation_priority": "Media"},
        {"id": "SR-07", "domain": "Infraestructura de Computo - Servers", "title": "SSL/TLS renegotiation DoS vulnerability", "affected_hosts": "1", "severity": "Medio", "description": "Renegociacion TLS habilitada.", "recommendation": "Deshabilitar renegociacion insegura.", "evidence": "Soporte a renegociacion vulnerable.", "source_tool": "Nessus", "status": "Pendiente", "remediation_priority": "Media"},
        {"id": "SR-09", "domain": "Infraestructura de Computo - Servers", "title": "Web Server Allows Password Auto-Completion", "affected_hosts": "1", "severity": "Bajo", "description": "Autocomplete de password habilitado.", "recommendation": "Deshabilitar autocompletado en formularios sensibles.", "evidence": "Formulario web permite password autocomplete.", "source_tool": "MonApps", "status": "Pendiente", "remediation_priority": "Baja"},
        {"id": "SW-01", "domain": "Infraestructura de Red - Switches", "title": "Generic HTTP Directory Traversal (Web Root)", "affected_hosts": "2", "severity": "Alto", "description": "Traversal sobre interfaz web de switches.", "recommendation": "Aplicar firmware y restringir acceso.", "evidence": "Active check confirma traversal.", "source_tool": "Nessus", "status": "Pendiente", "remediation_priority": "Alta"},
        {"id": "SW-02", "domain": "Infraestructura de Red - Switches", "title": "Vulnerable cipher suites for HTTPS", "affected_hosts": "2", "severity": "Alto", "description": "HTTPS inseguro en gestion de switches.", "recommendation": "Actualizar baseline TLS.", "evidence": "Suites vulnerables detectadas.", "source_tool": "Nessus", "status": "Pendiente", "remediation_priority": "Alta"},
        {"id": "SW-03", "domain": "Infraestructura de Red - Switches", "title": "Weak cipher suites", "affected_hosts": "2", "severity": "Medio", "description": "Cipher suites debiles habilitadas.", "recommendation": "Eliminar suites legadas.", "evidence": "Enumeracion TLS debil.", "source_tool": "OpenVAS", "status": "Pendiente", "remediation_priority": "Media"},
        {"id": "SW-04", "domain": "Infraestructura de Red - Switches", "title": "RSA keys less than 1024 bits", "affected_hosts": "2", "severity": "Medio", "description": "Claves RSA insuficientes.", "recommendation": "Renovar certificados y claves.", "evidence": "Cadena TLS con claves RSA debiles.", "source_tool": "Nessus", "status": "Pendiente", "remediation_priority": "Media"},
        {"id": "SW-06", "domain": "Infraestructura de Red - Switches", "title": "Weak Host Key Algorithms (SSH)", "affected_hosts": "1", "severity": "Medio", "description": "Host keys SSH debiles.", "recommendation": "Actualizar algoritmos SSH.", "evidence": "Enumeracion SSH insegura.", "source_tool": "OpenVAS", "status": "Pendiente", "remediation_priority": "Media"},
        {"id": "AP01", "domain": "Infraestructura de Red - WIFI", "title": "Deprecated SSH-1 protocol detection", "affected_hosts": "1", "severity": "Alto", "description": "SSH-1 habilitado en AP.", "recommendation": "Migrar a SSH moderno.", "evidence": "Servicio responde protocolo heredado.", "source_tool": "OpenVAS", "status": "Pendiente", "remediation_priority": "Alta"},
        {"id": "AP02", "domain": "Infraestructura de Red - WIFI", "title": "Weak Host Key Algorithms", "affected_hosts": "1", "severity": "Medio", "description": "Host keys SSH debiles en AP.", "recommendation": "Actualizar firmware y baseline SSH.", "evidence": "Algoritmos inseguros enumerados.", "source_tool": "OpenVAS", "status": "Pendiente", "remediation_priority": "Media"},
        {"id": "AP03", "domain": "Infraestructura de Red - WIFI", "title": "Weak KEX algorithms supported", "affected_hosts": "1", "severity": "Medio", "description": "KEX debiles en AP.", "recommendation": "Eliminar KEX heredados.", "evidence": "Enumeracion SSH con KEX debiles.", "source_tool": "OpenVAS", "status": "Pendiente", "remediation_priority": "Media"},
        {"id": "AP05", "domain": "Infraestructura de Red - WIFI", "title": "Deprecated TLSv1.0/TLSv1.1 detection", "affected_hosts": "1", "severity": "Medio", "description": "TLS obsoleto en AP.", "recommendation": "Deshabilitar TLS 1.0/1.1.", "evidence": "Handshake permite protocolos heredados.", "source_tool": "Nessus", "status": "Pendiente", "remediation_priority": "Media"},
        {"id": "DK-01", "domain": "Infraestructura de Computo - Desktops", "title": "Deprecated SSH-1 protocol detection", "affected_hosts": "1", "severity": "Alto", "description": "SSH heredado presente en estaciones puntuales.", "recommendation": "Deshabilitar protocolo legado.", "evidence": "Respuesta SSH-1 detectada.", "source_tool": "OpenVAS", "status": "Pendiente", "remediation_priority": "Alta"},
        {"id": "DK-02", "domain": "Infraestructura de Computo - Desktops", "title": "DCE/RPC and MSRPC services enumeration reporting", "affected_hosts": "40", "severity": "Medio", "description": "Servicios RPC expuestos en estaciones.", "recommendation": "Reducir exposicion y endurecer servicios.", "evidence": "Enumeracion RPC exitosa.", "source_tool": "MonInfra", "status": "Pendiente", "remediation_priority": "Media"},
        {"id": "DK-03", "domain": "Infraestructura de Computo - Desktops", "title": "Deprecated TLSv1.0/TLSv1.1 protocol detection", "affected_hosts": "42", "severity": "Medio", "description": "TLS obsoleto en estaciones.", "recommendation": "Forzar configuracion moderna.", "evidence": "Protocolos heredados permitidos.", "source_tool": "Nessus", "status": "Pendiente", "remediation_priority": "Media"},
        {"id": "DK-04", "domain": "Infraestructura de Computo - Desktops", "title": "Weak Host Key Algorithms (SSH)", "affected_hosts": "2", "severity": "Medio", "description": "Host keys SSH debiles.", "recommendation": "Actualizar hardening SSH.", "evidence": "Enumeracion SSH insegura.", "source_tool": "OpenVAS", "status": "Pendiente", "remediation_priority": "Media"},
        {"id": "OT-01", "domain": "Infraestructura OT/IoT", "title": "Firebird default credentials", "affected_hosts": "1", "severity": "Critico", "description": "Credenciales por defecto sobre servicio industrial.", "recommendation": "Rotar credenciales y segmentar acceso.", "evidence": "Login exitoso con credenciales default.", "source_tool": "Nessus", "status": "Pendiente", "remediation_priority": "Critica"},
        {"id": "OT-02", "domain": "Infraestructura OT/IoT", "title": "Vulnerable cipher suites for HTTPS", "affected_hosts": "5", "severity": "Alto", "description": "HTTPS inseguro en activos OT.", "recommendation": "Actualizar configuracion criptografica.", "evidence": "Enumeracion TLS vulnerable.", "source_tool": "OpenVAS", "status": "Pendiente", "remediation_priority": "Alta"},
        {"id": "OT-03", "domain": "Infraestructura OT/IoT", "title": "Deprecated SSLv2 and SSLv3 protocol detection", "affected_hosts": "4", "severity": "Medio", "description": "Protocolos obsoletos habilitados.", "recommendation": "Deshabilitar SSLv2/SSLv3.", "evidence": "Handshake permite SSL heredado.", "source_tool": "Nessus", "status": "Pendiente", "remediation_priority": "Media"},
        {"id": "OT-04", "domain": "Infraestructura OT/IoT", "title": "Weak cipher suites", "affected_hosts": "4", "severity": "Medio", "description": "Suites debiles en activos OT.", "recommendation": "Ajustar configuracion TLS.", "evidence": "Cipher suites inseguras enumeradas.", "source_tool": "OpenVAS", "status": "Pendiente", "remediation_priority": "Media"}
    ],
    "technical_evidence": [
        {"finding_id": "SR-01", "asset": "esxi-prd-01", "severity": "Critico", "source_tool": "Nessus", "evidence": "Operating System End of Life Detection", "impact": "Mayor exposicion a fallas sin soporte", "recommendation": "Actualizar o migrar plataforma", "status": "Pendiente"},
        {"finding_id": "OT-01", "asset": "ot-firebird-01", "severity": "Critico", "source_tool": "Nessus", "evidence": "Credenciales default validadas", "impact": "Compromiso administrativo inmediato", "recommendation": "Rotar credenciales y restringir acceso", "status": "En remediacion"},
        {"finding_id": "WE01", "asset": "portal.jockeysalud.pe", "severity": "Alto", "source_tool": "MonApps", "evidence": "Token embebido en JavaScript publico", "impact": "Posible abuso de sesiones o APIs", "recommendation": "Rotar secreto y sanear recurso", "status": "Pendiente"},
        {"finding_id": "SW-02", "asset": "switch-core-02", "severity": "Alto", "source_tool": "Nessus", "evidence": "Cipher suites vulnerables en interfaz HTTPS", "impact": "Debilitamiento criptografico de administracion", "recommendation": "Aplicar baseline TLS actualizado", "status": "Pendiente"},
        {"finding_id": "AP01", "asset": "ap-core-01", "severity": "Alto", "source_tool": "OpenVAS", "evidence": "SSH-1 habilitado", "impact": "Riesgo elevado de acceso inseguro", "recommendation": "Deshabilitar protocolo heredado", "status": "Pendiente"},
    ],
    "operational_timeline": [
        {"date": "2026-06-20", "event": "Inicio del corte semanal y consolidacion de hallazgos", "category": "Monitoreo", "impact": "Se habilita priorizacion inicial del backlog"},
        {"date": "2026-06-21", "event": "Validacion de credenciales default en OT", "category": "Validacion", "impact": "Se abre caso critico con riesgo administrativo"},
        {"date": "2026-06-22", "event": "Revision criptografica en switches y APs", "category": "Hardening", "impact": "Se priorizan hallazgos altos y medios sobre gestion"},
        {"date": "2026-06-24", "event": "Correlacion de hallazgos web externos", "category": "Analisis", "impact": "Se confirma exposicion de token y TLS vulnerable"},
        {"date": "2026-06-25", "event": "Cierre parcial sobre desktops", "category": "Remediacion", "impact": "Se reduce parte del backlog medio"},
        {"date": "2026-06-26", "event": "Consolidacion ejecutiva del informe", "category": "Reporte", "impact": "Se emite salida semanal para seguimiento"},
    ],
    "actions_worked": [
        "Consolidacion de hallazgos sobre activos web, red, servidores y OT priorizados.",
        "Validacion tecnica de credenciales default, TLS debil y plataformas fuera de soporte.",
        "Seguimiento del backlog de remediacion sobre switching, APs y estaciones.",
        "Correlacion de evidencia tecnica para priorizacion de criticidades abiertas.",
        "Construccion del resumen ejecutivo semanal para seguimiento con el cliente.",
    ],
    "automation_suggestions": [
        {"name": "Alerta automatica de credenciales por defecto", "domain": "OT/IoT", "priority": "Alta", "benefit": "Reducir tiempo de deteccion de riesgos criticos administrativos", "status": "Propuesta"},
        {"name": "Monitoreo continuo de cipher suites HTTPS", "domain": "Web Externo / Switches", "priority": "Alta", "benefit": "Detectar regresiones criptograficas temprano", "status": "Propuesta"},
        {"name": "Inventario automatizado de activos EOL", "domain": "Servers", "priority": "Alta", "benefit": "Priorizar plataformas fuera de soporte", "status": "Propuesta"},
        {"name": "Correlacion de TLS heredado en APs", "domain": "WIFI", "priority": "Media", "benefit": "Reducir deuda tecnica de seguridad inalambrica", "status": "Propuesta"},
        {"name": "Seguimiento automatizado de RPC expuesto", "domain": "Desktops", "priority": "Media", "benefit": "Reducir superficie lateral en estaciones", "status": "Propuesta"},
    ],
    "next_actions": [
        {"action": "Actualizar ESXi fuera de soporte", "owner": "Cliente", "priority": "Critica", "due_date": "2026-07-03", "dependency": "Ventana de mantenimiento"},
        {"action": "Rotar credenciales por defecto en activo OT", "owner": "Cliente", "priority": "Critica", "due_date": "2026-06-27", "dependency": "Aprobacion de acceso"},
        {"action": "Endurecer cipher suites en switches core", "owner": "Cliente", "priority": "Alta", "due_date": "2026-07-01", "dependency": "Validacion de firmware"},
        {"action": "Eliminar token expuesto del portal externo", "owner": "Cliente", "priority": "Alta", "due_date": "2026-06-27", "dependency": "Cambio aplicativo"},
        {"action": "Deshabilitar TLS heredado en APs", "owner": "XOC / Cliente", "priority": "Media", "due_date": "2026-07-02", "dependency": "Revision de compatibilidad"},
    ],
    "results": {
        "summary": "Se consolido una vista semanal ejecutiva y tecnica de la postura de seguridad de Jockey Salud, con foco en criticidades abiertas, riesgos por dominio y acciones inmediatas.",
        "highlights": [
            "Se identifico mejora parcial en backlog de desktops.",
            "Se mantuvo trazabilidad entre hallazgos, casos y SLA.",
            "Persisten riesgos altos y criticos en servidores y OT/IoT.",
        ],
    },
    "pending_findings": [
        "Actualizacion del hipervisor fuera de soporte.",
        "Rotacion de credenciales por defecto en activo OT.",
        "Hardening criptografico de switches y APs.",
        "Remocion de token expuesto en recurso web publico.",
    ],
    "security_news": [
        {"title": "Campanas de abuso de credenciales contra portales corporativos", "date": "2026-06-25", "source": "XOC Threat Intel", "summary": "Se mantienen campanas orientadas al robo de credenciales y abuso de sesiones sobre portales expuestos a Internet.", "links": ["https://example.com/news/credential-theft"]},
        {"title": "Nueva alerta de seguridad sobre componentes perimetrales heredados", "date": "2026-06-24", "source": "Vendor Advisory", "summary": "Fabricantes publicaron alertas y parches para componentes perimetrales con exposicion remota no autenticada.", "links": ["https://example.com/news/perimeter-rce"]},
    ],
    "annex": [
        {"reference": "Consolidado MonVulE/MonVulC", "detail": "Backlog tecnico consolidado por dominio y severidad."},
        {"reference": "Bitacora operativa XOC", "detail": "Seguimiento de acciones, owners y ventanas de remediacion."},
        {"reference": "Evidencia Wazuh, Zabbix y Nessus", "detail": "Soporte tecnico usado para validacion del corte semanal."},
    ],
}


def deep_merge(base: Any, override: Any) -> Any:
    if isinstance(base, dict) and isinstance(override, dict):
        merged = deepcopy(base)
        for key, value in override.items():
            merged[key] = deep_merge(merged[key], value) if key in merged else deepcopy(value)
        return merged
    if override is None:
        return deepcopy(base)
    return deepcopy(override)


def ensure_directories() -> None:
    for path in (DATA_DIR, TEMPLATES_DIR, OUTPUT_DIR, REPORTS_ROOT_DIR):
        path.mkdir(parents=True, exist_ok=True)
    print("Directorios OK")


def create_mock_data_if_missing() -> None:
    if not MOCK_DATA_PATH.exists():
        MOCK_DATA_PATH.write_text(json.dumps(DEFAULT_MOCK_DATA, indent=2, ensure_ascii=False), encoding="utf-8")


def load_mock_data() -> dict[str, Any]:
    data = json.loads(MOCK_DATA_PATH.read_text(encoding="utf-8"))
    merged = deep_merge(DEFAULT_MOCK_DATA, data)
    if merged != data:
        MOCK_DATA_PATH.write_text(json.dumps(merged, indent=2, ensure_ascii=False), encoding="utf-8")
    return merged


def build_report_output_path(data: dict[str, Any]) -> Path:
    return REPORTS_ROOT_DIR / data["tenant"]["id"] / data["report"]["id"] / "v1-generated.docx"


def load_report_data() -> tuple[dict[str, Any], str]:
    use_database = os.environ.get("USE_DATABASE", "false").strip().lower() == "true"
    mock_data = load_mock_data()
    if not use_database:
        print("Fuente de datos: mock")
        return mock_data, "mock"

    print("Fuente de datos: BD")
    try:
        report_data = build_report_data_from_database(mock_data)
        DATABASE_REPORT_DATA_PATH.write_text(json.dumps(report_data, indent=2, ensure_ascii=False), encoding="utf-8")
        print("Conexion BD OK")
        return report_data, "database"
    except Exception as exc:
        print(f"WARNING: conexion BD o carga de datos fallo, fallback a mock. Detalle: {exc}")
        return mock_data, "mock_fallback"


def calculate_derived_metrics(data: dict[str, Any]) -> dict[str, Any]:
    previous = data.get("severity_summary", {}).get("previous", {})
    current = data.get("severity_summary", {}).get("current", {})
    severity_delta = {
        key: current.get(key, 0) - previous.get(key, 0)
        for key in ("critical", "high", "medium", "low", "informational")
    }

    sla_source = data.get("sla_summary") or {}
    if not sla_source or "total_cases" not in sla_source:
        counts = Counter(case.get("sla_status", "") for case in data.get("cases", []))
        sla_source = {
            "total_cases": len(data.get("cases", [])),
            "within_sla": counts.get("Dentro SLA", 0),
            "at_risk": counts.get("En riesgo", 0),
            "breached": counts.get("Vencido", 0),
            "average_elapsed_hours": round(sum(case.get("elapsed_hours", 0) for case in data.get("cases", [])) / max(1, len(data.get("cases", []))), 2),
        }

    open_cases = sum(1 for case in data.get("cases", []) if case.get("status") != "Resuelto")
    resolved_cases = sum(1 for case in data.get("cases", []) if case.get("status") == "Resuelto")
    case_severity_counts = Counter(case.get("severity", "") for case in data.get("cases", []))
    critical_assets = sum(1 for asset in data.get("top_assets", []) if asset.get("critical", 0) > 0)

    posture = data.setdefault("security_posture", {})
    posture.setdefault("global_score", round(sum(item.get("score", 0) for item in data.get("domain_scores", [])) / max(1, len(data.get("domain_scores", []))), 2))
    posture.setdefault("risk_level", "Medio")
    posture.setdefault("trend", "Estable")
    posture.setdefault("critical_open", case_severity_counts.get("Critico", 0))
    posture.setdefault("high_open", case_severity_counts.get("Alto", 0))
    posture.setdefault("sla_at_risk", sla_source.get("at_risk", 0))
    posture.setdefault("mttr_hours", sla_source.get("average_elapsed_hours", 0))
    posture["top_risks"] = posture.get("main_risks") or posture.get("top_risks") or []

    data["severity_delta"] = {
        **severity_delta,
        "risk_trend": "Mejora" if severity_delta["critical"] < 0 else "Empeora" if severity_delta["critical"] > 0 or severity_delta["high"] > 0 else "Estable",
        "narrative": (
            f"La semana actual registra una variacion de {severity_delta['critical']} criticos, {severity_delta['high']} altos, "
            f"{severity_delta['medium']} medios, {severity_delta['low']} bajos y {severity_delta['informational']} informativos respecto del corte anterior."
        ),
    }
    data["sla_summary"] = sla_source
    data["kpis"] = {
        "Casos abiertos": open_cases,
        "Casos resueltos": resolved_cases,
        "Hallazgos nuevos": sum(1 for value in severity_delta.values() if value > 0),
        "Hallazgos resueltos": sum(1 for finding in data.get("findings", []) if finding.get("status") == "Resuelto"),
        "SLA dentro de plazo": sla_source.get("within_sla", 0),
        "SLA en riesgo": sla_source.get("at_risk", 0),
        "SLA vencido": sla_source.get("breached", 0),
        "Activos criticos afectados": critical_assets,
    }
    data["report"]["generated_at"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    print("Metricas calculadas OK")
    return data


def _save_chart(fig: Any, path: Path) -> Path:
    fig.tight_layout()
    fig.savefig(path, dpi=200, bbox_inches="tight")
    plt.close(fig)
    return path


def generate_charts(data: dict[str, Any]) -> dict[str, Path]:
    previous = data["severity_summary"]["previous"]
    current = data["severity_summary"]["current"]
    labels = ["Critico", "Alto", "Medio", "Bajo", "Informativo"]
    prev_values = [previous["critical"], previous["high"], previous["medium"], previous["low"], previous["informational"]]
    curr_values = [current["critical"], current["high"], current["medium"], current["low"], current["informational"]]

    fig, ax = plt.subplots(figsize=(9, 4.8))
    x = list(range(len(labels)))
    ax.bar([i - 0.2 for i in x], prev_values, width=0.4, color="#94a3b8", label="Semana anterior")
    ax.bar([i + 0.2 for i in x], curr_values, width=0.4, color="#1d4ed8", label="Semana actual")
    ax.set_title("Comparativo semanal por severidad")
    ax.set_xticks(x)
    ax.set_xticklabels(labels)
    ax.grid(axis="y", linestyle="--", alpha=0.3)
    ax.legend()
    _save_chart(fig, SEVERITY_CHART_PATH)

    fig, ax = plt.subplots(figsize=(10, 5))
    domain_labels = [item["domain"] for item in data.get("domain_scores", [])]
    domain_values = [item["score"] for item in data.get("domain_scores", [])]
    ax.barh(domain_labels, domain_values, color="#0f766e")
    ax.set_xlim(0, 100)
    ax.set_title("Score por dominio")
    ax.grid(axis="x", linestyle="--", alpha=0.3)
    _save_chart(fig, DOMAIN_SCORE_CHART_PATH)

    fig, ax = plt.subplots(figsize=(7, 4.2))
    case_counts = Counter(case.get("severity", "") for case in data.get("cases", []))
    severity_order = ["Critico", "Alto", "Medio", "Bajo"]
    ax.bar(severity_order, [case_counts.get(key, 0) for key in severity_order], color=["#b91c1c", "#ea580c", "#d97706", "#facc15"])
    ax.set_title("Casos por severidad")
    ax.grid(axis="y", linestyle="--", alpha=0.3)
    _save_chart(fig, CASES_BY_SEVERITY_CHART_PATH)

    fig, ax = plt.subplots(figsize=(6.6, 4.2))
    values = [data["sla_summary"].get("within_sla", 0), data["sla_summary"].get("at_risk", 0), data["sla_summary"].get("breached", 0)]
    ax.pie(values if sum(values) > 0 else [1], labels=["Dentro SLA", "En riesgo", "Vencido"] if sum(values) > 0 else ["Sin datos"], autopct="%1.0f%%" if sum(values) > 0 else None, colors=["#16a34a", "#f59e0b", "#dc2626"][: 3 if sum(values) > 0 else 1])
    ax.set_title("Estado SLA")
    _save_chart(fig, SLA_STATUS_CHART_PATH)

    print("Charts generados OK")
    return {
        "severity_comparison_chart": SEVERITY_CHART_PATH,
        "domain_score_chart": DOMAIN_SCORE_CHART_PATH,
        "cases_by_severity_chart": CASES_BY_SEVERITY_CHART_PATH,
        "sla_status_chart": SLA_STATUS_CHART_PATH,
    }


def build_docx_context(data: dict[str, Any], ai_sections: dict[str, str], charts: dict[str, Path]) -> dict[str, Any]:
    context = deepcopy(data)
    context["ai_sections"] = ai_sections
    context["charts"] = charts
    context["generated_at_display"] = datetime.strptime(data["report"]["generated_at"], "%Y-%m-%d %H:%M:%S").strftime("%d/%m/%Y %H:%M")
    return context


def configure_section(section, *, landscape: bool = False) -> None:
    section.orientation = WD_ORIENT.LANDSCAPE if landscape else WD_ORIENT.PORTRAIT
    if landscape:
        section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)


def set_document_defaults(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(10.5)
    for section in document.sections:
        configure_section(section)


def clear_template_body_after_cover(document: Document, marker_text: str = "Contenido") -> None:
    marker = next((paragraph for paragraph in document.paragraphs if paragraph.text.strip() == marker_text), None)
    if marker is None:
        return
    body = document.element.body
    marker_element = marker._p
    removing = False
    for child in list(body):
        if child is marker_element:
            removing = True
        if removing and child.tag != qn("w:sectPr"):
            body.remove(child)


def render_cover_placeholders_if_present(document: Document, context: dict[str, Any]) -> None:
    replacements = {
        "tenant.name": context["tenant"]["name"],
        "tenant.id": context["tenant"]["id"],
        "tenant.tenant_id": context["tenant"]["tenant_id"],
        "report.id": context["report"]["id"],
        "report.title": context["report"]["title"],
        "report.period": context["report"]["period"],
        "report.prepared_by": context["report"]["prepared_by"],
        "report.service": context["report"]["service"],
    }
    for paragraph in document.paragraphs:
        for key, value in replacements.items():
            for token in (f"{{{{{key}}}}}", f"{{{{ {key} }}}}"):
                if token in paragraph.text:
                    paragraph.text = paragraph.text.replace(token, str(value))


def add_page_break_after_cover(document: Document) -> None:
    section = document.add_section(WD_SECTION.NEW_PAGE)
    configure_section(section)


def start_landscape_section(document: Document) -> None:
    section = document.add_section(WD_SECTION.NEW_PAGE)
    configure_section(section, landscape=True)


def start_portrait_section(document: Document) -> None:
    section = document.add_section(WD_SECTION.NEW_PAGE)
    configure_section(section)


def add_report_title(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(15, 23, 42)
    paragraph.paragraph_format.space_after = Pt(8)


def add_section_heading(document: Document, number: int, title: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(f"{number}. {title}")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(17, 24, 39)


def add_body_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_after = Pt(6)


def add_bullet_list(document: Document, items: list[str]) -> None:
    for item in items:
        try:
            paragraph = document.add_paragraph(style="List Bullet")
        except KeyError:
            paragraph = document.add_paragraph()
            paragraph.add_run("- ")
        paragraph.add_run(item)
        paragraph.paragraph_format.space_after = Pt(2)


def style_table(table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(9)
                    if row_index == 0:
                        run.bold = True
            if row_index == 0:
                tc_pr = cell._tc.get_or_add_tcPr()
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), "DCE6F1")
                tc_pr.append(shading)


def add_simple_table(document: Document, headers: list[str], rows: list[list[Any]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = str(header)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = str(value)
    style_table(table)


def add_chart(document: Document, image_path: Path, width_inches: float) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(image_path), width=Inches(width_inches))


def add_manual_table_of_contents(document: Document, sections: list[str]) -> None:
    for section in sections:
        add_body_paragraph(document, section)


def append_report_body_to_existing_template(template_path: Path, output_path: Path, context: dict[str, Any]) -> Path:
    document = Document(str(template_path))
    set_document_defaults(document)
    render_cover_placeholders_if_present(document, context)
    clear_template_body_after_cover(document)
    add_page_break_after_cover(document)
    print("Portada cargada desde plantilla existente")

    add_report_title(document, context["report"]["title"])

    add_section_heading(document, 1, "Tabla de contenido")
    add_manual_table_of_contents(
        document,
        [
            "1. Tabla de contenido manual/simple",
            "2. Datos generales",
            "3. Herramientas utilizadas",
            "4. Resumen ejecutivo",
            "5. Security Posture Overview",
            "6. KPIs operativos",
            "7. Analisis semanal de vulnerabilidades",
            "8. Comparativo semana anterior vs actual",
            "9. Grafico de severidades",
            "10. Score por dominio",
            "11. Cases / Incidentes XOC",
            "12. SLA / MTTR",
            "13. MITRE ATT&CK Mapping",
            "14. Top activos afectados",
            "15. Seguridad por dominio",
            "16. Evidencia tecnica critica/alta",
            "17. Timeline operativo semanal",
            "18. Acciones trabajadas",
            "19. Automatizaciones recomendadas",
            "20. Proximas acciones",
            "21. Resultados obtenidos",
            "22. Hallazgos pendientes",
            "23. Noticias de seguridad",
            "24. Conclusion de riesgo",
            "25. Notas del analista",
            "26. Anexo tecnico",
        ],
    )

    add_section_heading(document, 2, "Datos generales")
    add_simple_table(document, ["Campo", "Valor"], [["Cliente", context["tenant"]["name"]], ["Tenant ID", context["tenant"]["tenant_id"]], ["Servicio", context["report"]["service"]], ["Periodo", context["report"]["period"]], ["Fecha de generacion", context["generated_at_display"]], ["Preparado por", context["report"]["prepared_by"]], ["Alcance", context["report"]["scope"]]])

    add_section_heading(document, 3, "Herramientas utilizadas")
    add_bullet_list(document, context["tools"])

    start_portrait_section(document)
    add_section_heading(document, 4, "Resumen ejecutivo")
    add_body_paragraph(document, context["ai_sections"]["executive_summary"])

    add_section_heading(document, 5, "Security Posture Overview")
    add_simple_table(document, ["Indicador", "Valor"], [["Score global", context["security_posture"]["global_score"]], ["Nivel de riesgo", context["security_posture"]["risk_level"]], ["Tendencia", context["security_posture"]["trend"]], ["Casos criticos abiertos", context["security_posture"]["critical_open"]], ["Casos altos abiertos", context["security_posture"]["high_open"]], ["SLA en riesgo", context["security_posture"]["sla_at_risk"]], ["MTTR", f"{context['security_posture']['mttr_hours']} horas"]])
    add_body_paragraph(document, context["ai_sections"]["security_posture_narrative"])
    add_bullet_list(document, context["security_posture"].get("top_risks", []))

    add_section_heading(document, 6, "KPIs operativos")
    add_simple_table(document, ["KPI", "Valor"], [[key, value] for key, value in context["kpis"].items()])

    start_portrait_section(document)
    add_section_heading(document, 7, "Analisis semanal de vulnerabilidades")
    add_body_paragraph(document, context["ai_sections"]["severity_analysis"])

    add_section_heading(document, 8, "Comparativo semana anterior vs actual")
    previous = context["severity_summary"]["previous"]
    current = context["severity_summary"]["current"]
    add_simple_table(document, ["Severidad", "Semana anterior", "Semana actual", "Delta"], [["Critico", previous["critical"], current["critical"], context["severity_delta"]["critical"]], ["Alto", previous["high"], current["high"], context["severity_delta"]["high"]], ["Medio", previous["medium"], current["medium"], context["severity_delta"]["medium"]], ["Bajo", previous["low"], current["low"], context["severity_delta"]["low"]], ["Informativo", previous["informational"], current["informational"], context["severity_delta"]["informational"]]])

    add_section_heading(document, 9, "Grafico de severidades")
    add_chart(document, context["charts"]["severity_comparison_chart"], 6.4)

    add_section_heading(document, 10, "Score por dominio")
    add_chart(document, context["charts"]["domain_score_chart"], 6.6)
    add_simple_table(document, ["Dominio", "Score", "Riesgo", "Tendencia", "Comentario"], [[item["domain"], item["score"], item["risk"], item["trend"], item["comment"]] for item in context["domain_scores"]])

    start_landscape_section(document)
    add_section_heading(document, 11, "Cases / Incidentes XOC")
    add_body_paragraph(document, context["ai_sections"]["cases_summary"])
    add_chart(document, context["charts"]["cases_by_severity_chart"], 5.6)
    add_simple_table(document, ["Caso", "Severidad", "Estado", "Dominio", "Activo", "SLA", "Responsable", "Proxima accion"], [[case["id"], case["severity"], case["status"], case["domain"], case["asset"], case["sla_status"], case["owner"], case["next_action"]] for case in context["cases"]])

    add_section_heading(document, 12, "SLA / MTTR")
    add_body_paragraph(document, context["ai_sections"]["sla_summary_narrative"])
    add_chart(document, context["charts"]["sla_status_chart"], 4.6)
    add_simple_table(document, ["Indicador", "Valor"], [["Total de casos", context["sla_summary"].get("total_cases", 0)], ["Dentro SLA", context["sla_summary"].get("within_sla", 0)], ["En riesgo", context["sla_summary"].get("at_risk", 0)], ["Vencido", context["sla_summary"].get("breached", 0)], ["Promedio horas transcurridas", context["sla_summary"].get("average_elapsed_hours", 0)]])

    add_section_heading(document, 13, "MITRE ATT&CK Mapping")
    add_body_paragraph(document, context["ai_sections"]["mitre_summary"])
    add_simple_table(document, ["Tactica", "Tecnica", "Evidencia", "Dominio", "Severidad", "Estado"], [[item["tactic"], item["technique"], item["evidence"], item["domain"], item["severity"], item["status"]] for item in context["mitre_mapping"]])

    add_section_heading(document, 14, "Top activos afectados")
    add_body_paragraph(document, context["ai_sections"]["top_assets_summary"])
    add_simple_table(document, ["Activo", "Tipo", "Dominio", "Score de riesgo", "Criticos", "Altos", "Medios", "Bajos", "Recomendacion"], [[item["asset"], item["type"], item["domain"], item["risk_score"], item["critical"], item["high"], item["medium"], item["low"], item["recommendation"]] for item in context["top_assets"]])

    start_portrait_section(document)
    add_section_heading(document, 15, "Seguridad por dominio")
    for domain in context["domains"]:
        add_body_paragraph(document, f"{domain['name']}: Score {domain['score']}/100, riesgo {domain['risk']}. {domain['summary']}")

    start_landscape_section(document)
    add_section_heading(document, 16, "Evidencia tecnica critica/alta")
    add_body_paragraph(document, context["ai_sections"]["technical_evidence_summary"])
    add_simple_table(document, ["ID finding", "Activo", "Severidad", "Herramienta fuente", "Evidencia", "Impacto", "Recomendacion", "Estado"], [[item["finding_id"], item["asset"], item["severity"], item["source_tool"], item["evidence"], item["impact"], item["recommendation"], item["status"]] for item in context["technical_evidence"]])

    add_section_heading(document, 17, "Timeline operativo semanal")
    add_body_paragraph(document, context["ai_sections"]["operational_timeline_summary"])
    add_simple_table(document, ["Fecha", "Evento", "Categoria", "Impacto"], [[item["date"], item["event"], item["category"], item["impact"]] for item in context["operational_timeline"]])

    add_section_heading(document, 18, "Acciones trabajadas")
    add_bullet_list(document, context["actions_worked"])

    add_section_heading(document, 19, "Automatizaciones recomendadas")
    add_body_paragraph(document, context["ai_sections"]["automation_recommendations_summary"])
    add_simple_table(document, ["Nombre", "Dominio", "Prioridad", "Beneficio", "Estado"], [[item["name"], item["domain"], item["priority"], item["benefit"], item["status"]] for item in context["automation_suggestions"]])

    add_section_heading(document, 20, "Proximas acciones")
    add_body_paragraph(document, context["ai_sections"]["next_actions_summary"])
    add_simple_table(document, ["Accion", "Owner", "Prioridad", "Fecha objetivo", "Dependencia"], [[item["action"], item["owner"], item["priority"], item["due_date"], item["dependency"]] for item in context["next_actions"]])

    add_section_heading(document, 21, "Resultados obtenidos")
    results = context["results"]
    if isinstance(results, dict):
        add_body_paragraph(document, results.get("summary", ""))
        add_bullet_list(document, results.get("highlights", []))
    else:
        add_bullet_list(document, results)

    add_section_heading(document, 22, "Hallazgos pendientes")
    add_bullet_list(document, context["pending_findings"])

    add_section_heading(document, 23, "Noticias de seguridad")
    add_simple_table(document, ["Noticia", "Fecha", "Fuente", "Resumen"], [[item["title"], item["date"], item["source"], item["summary"]] for item in context["security_news"]])

    add_section_heading(document, 24, "Conclusion de riesgo")
    add_body_paragraph(document, context["ai_sections"]["risk_conclusion"])

    add_section_heading(document, 25, "Notas del analista")
    add_body_paragraph(document, context["ai_sections"]["analyst_notes"])

    add_section_heading(document, 26, "Anexo tecnico")
    add_simple_table(document, ["Referencia", "Detalle"], [[item["reference"], item["detail"]] for item in context["annex"]])

    document.save(str(output_path))
    print("Cuerpo agregado despues de portada")
    return output_path


def validate_docx(path: Path) -> tuple[bool, list[str]]:
    messages: list[str] = []
    if not path.exists():
        return False, ["El DOCX final no existe"]
    if path.stat().st_size <= 20 * 1024:
        return False, ["El DOCX final pesa 20KB o menos"]
    with zipfile.ZipFile(path, "r") as archive:
        if "word/document.xml" not in archive.namelist():
            return False, ["No existe word/document.xml"]
        document_xml = archive.read("word/document.xml")
    messages.append("Validacion zipfile OK")
    etree.fromstring(document_xml)
    messages.append("Validacion XML OK")
    return True, messages


def sha256_file(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def main() -> None:
    ensure_directories()
    create_mock_data_if_missing()
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"No existe la plantilla base: {TEMPLATE_PATH}")
    template_hash_before = sha256_file(TEMPLATE_PATH)

    report_data, _source = load_report_data()
    report_data = calculate_derived_metrics(report_data)
    charts = generate_charts(report_data)
    ai_sections, ai_source = generate_ai_sections(report_data, AZURE_SECTIONS_OUTPUT_PATH)
    if ai_source in {"azure", "azure_openai"}:
        print("Azure usado OK")
    else:
        print("Fallback local usado OK")
    context = build_docx_context(report_data, ai_sections, charts)

    output_path = build_report_output_path(report_data)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    append_report_body_to_existing_template(TEMPLATE_PATH, output_path, context)
    valid, messages = validate_docx(output_path)
    if not valid:
        raise RuntimeError(" | ".join(messages))
    if sha256_file(TEMPLATE_PATH) != template_hash_before:
        raise RuntimeError("La plantilla original fue modificada y eso no esta permitido")
    print("DOCX generado OK")
    print("Validacion OK")
    for message in messages:
        print(message)
    print(f"Ruta final del documento: {output_path}")


if __name__ == "__main__":
    main()

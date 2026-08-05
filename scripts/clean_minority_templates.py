from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.table import Table
from copy import deepcopy


ROOT = Path(__file__).resolve().parents[1]
TEMPLATES = [
    ROOT / "Plantillas" / "TXDX_report_for_client-tenant MINORITY_REPORT_SEMANAL_dia_mes.docx",
    ROOT / "Plantillas" / "TXDX_report_for_admin_client-tenant MINORITY_REPORT_SEMANAL_dia_mes.docx",
]

BODY_START = "[[BODY_START]]"
DOMAIN_PROTO = "[[PROTO:DOMAIN_TABLE]]"
COMPARISON_PROTO = "[[PROTO:COMPARISON_TABLE]]"
COVER_CLIENT = "[[COVER_CLIENT]]"
COVER_PERIOD = "[[COVER_PERIOD]]"
COVER_PREPARED_BY = "[[COVER_PREPARED_BY]]"
FOOTER_CLIENT = "[[FOOTER_CLIENT]]"


def textbox_text(box) -> str:
    return "".join(node.text or "" for node in box.iter(qn("w:t")))


def set_textbox_text(box, value: str) -> None:
    nodes = list(box.iter(qn("w:t")))
    if not nodes:
        return
    nodes[0].text = value
    for node in nodes[1:]:
        node.text = ""


def remove_body_after_first_heading(document: Document) -> None:
    body = document._element.body
    children = list(body)
    start_index = None
    for index, child in enumerate(children):
        text = "".join(node.text or "" for node in child.iter(qn("w:t"))).strip()
        if BODY_START in text:
            start_index = index
            break
        if text == "Datos generales":
            start_index = index
            break
    if start_index is None:
        raise RuntimeError("No se encontró el inicio del cuerpo de ejemplo en la plantilla")
    for child in list(children[start_index:]):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def hide_paragraph(paragraph) -> None:
    for run in paragraph.runs:
        run.font.hidden = True
        run.font.size = None


def add_hidden_marker(document: Document, marker: str) -> None:
    paragraph = document.add_paragraph(marker, style="Body Text")
    hide_paragraph(paragraph)


def trim_table_rows(table: Table, keep_rows: int = 2) -> None:
    for row in list(table.rows[keep_rows:]):
        table._tbl.remove(row._tr)


def normalize_toc_cache(document: Document) -> None:
    body = document._element.body
    content_child = None
    for child in list(body):
        text = "".join(node.text or "" for node in child.iter(qn("w:t"))).strip()
        if text == "Contenido":
            content_child = child
            break
    if content_child is None:
        return
    sdt = content_child.getnext()
    if sdt is None or sdt.tag != qn("w:sdt"):
        return
    sdt_content = sdt.find(qn("w:sdtContent"))
    if sdt_content is None:
        return
    templates: list = []
    for paragraph in sdt_content.findall(qn("w:p")):
        templates.append(deepcopy(paragraph))
        if len(templates) == 3:
            break
    if not templates:
        return
    while len(templates) < 3:
        templates.append(deepcopy(templates[-1]))
    labels = ["1. Secciones dinámicas", "1.1. Subsección dinámica", "1.1.1. Nivel dinámico"]
    for child in list(sdt_content):
        sdt_content.remove(child)
    for template, label in zip(templates, labels):
        text_nodes = list(template.iter(qn("w:t")))
        if text_nodes:
            text_nodes[0].text = label
            for node in text_nodes[1:]:
                node.text = ""
        sdt_content.append(template)


def add_comparison_proto(document: Document) -> None:
    add_hidden_marker(document, COMPARISON_PROTO)
    table = document.add_table(rows=2, cols=3)
    headers = ["Severidad", "Semana Anterior", "Semana Actual"]
    for index, header in enumerate(headers):
        table.cell(0, index).text = header
    sample = ["ALTO", "0", "0"]
    for index, value in enumerate(sample):
        table.cell(1, index).text = value


def add_domain_proto(document: Document) -> None:
    add_hidden_marker(document, DOMAIN_PROTO)
    table = document.add_table(rows=2, cols=4)
    headers = ["ID", "Vulnerabilidades", "Host Afectados", "Severidad"]
    for index, header in enumerate(headers):
        table.cell(0, index).text = header
    sample = ["0000", "Hallazgo de ejemplo", "host.local", "ALTO"]
    for index, value in enumerate(sample):
        table.cell(1, index).text = value


def add_body_anchor(document: Document) -> None:
    paragraph = document.add_paragraph(style="Body Text")
    run = paragraph.add_run(BODY_START)
    run.font.hidden = True
    run.add_break(WD_BREAK.PAGE)


def normalize_cover_and_footer(document: Document) -> None:
    for box in document._element.xpath(".//w:txbxContent"):
        original = textbox_text(box).strip()
        if original == "Change for date":
            set_textbox_text(box, COVER_PERIOD)
        elif original == "Change for prepared for":
            set_textbox_text(box, COVER_CLIENT)
        elif original == "TXDXSECURE":
            set_textbox_text(box, COVER_PREPARED_BY)
        elif "JOCKEY" in original and "SALUD" in original:
            set_textbox_text(box, COVER_CLIENT)

    for paragraph in document.paragraphs[:45]:
        text = paragraph.text.strip()
        if "JOCKEY SALUD" in text and "TXDXSECURE" in text:
            paragraph.text = f"{COVER_CLIENT}\t{COVER_PREPARED_BY}"
        elif "Del 20 de junio al 26 de junio del 2026" in text or "Del 20 al 26 de junio del 2026" in text:
            paragraph.text = COVER_PERIOD

    for section in document.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        if not footer.paragraphs:
            footer.add_paragraph()
        footer.paragraphs[0].text = f"TxdxSecure\tMinority Report XOC\t{FOOTER_CLIENT}"
        for paragraph in footer.paragraphs[1:]:
            paragraph.text = ""


def clean_template(path: Path) -> None:
    document = Document(path)
    comparison_template = None
    domain_template = None
    for table in document.tables:
        cols = len(table.columns)
        if cols == 3 and comparison_template is None:
            comparison_template = deepcopy(table._tbl)
        elif cols == 4 and domain_template is None:
            domain_template = deepcopy(table._tbl)
    normalize_cover_and_footer(document)
    normalize_toc_cache(document)
    remove_body_after_first_heading(document)
    add_body_anchor(document)
    if comparison_template is not None:
        add_hidden_marker(document, COMPARISON_PROTO)
        marker = document.paragraphs[-1]
        marker._p.addnext(deepcopy(comparison_template))
        table = document.tables[-1]
        trim_table_rows(table)
        headers = ["Severidad", "Semana Anterior", "Semana Actual"]
        for index, header in enumerate(headers):
            table.cell(0, index).text = header
        sample = ["ALTO", "0", "0"]
        for index, value in enumerate(sample):
            table.cell(1, index).text = value
    else:
        add_comparison_proto(document)
    if domain_template is not None:
        add_hidden_marker(document, DOMAIN_PROTO)
        marker = document.paragraphs[-1]
        marker._p.addnext(deepcopy(domain_template))
        table = document.tables[-1]
        trim_table_rows(table)
        headers = ["ID", "Vulnerabilidades", "Host Afectados", "Severidad"]
        for index, header in enumerate(headers):
            table.cell(0, index).text = header
        sample = ["0000", "Hallazgo de ejemplo", "host.local", "ALTO"]
        for index, value in enumerate(sample):
            table.cell(1, index).text = value
    else:
        add_domain_proto(document)
    document.save(path)


def main() -> None:
    for template_path in TEMPLATES:
        clean_template(template_path)
        print(f"Cleaned template: {template_path}")


if __name__ == "__main__":
    main()
